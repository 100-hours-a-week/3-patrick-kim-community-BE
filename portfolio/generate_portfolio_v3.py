#!/usr/bin/env python3
"""PetStar Portfolio v3 — Black & White, technical diagram style"""

import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

# ── Palette: pure B&W ──
BK    = '#1A1A1A'
D1    = '#333333'
D2    = '#555555'
D3    = '#777777'
MD    = '#999999'
LT    = '#BBBBBB'
LN    = '#CCCCCC'
BG    = '#F2F2F2'
BG2   = '#E8E8E8'
WH    = '#FFFFFF'

FONT = 'Apple SD Gothic Neo'
plt.rcParams['font.family'] = FONT
plt.rcParams['axes.unicode_minus'] = False
OUT = os.path.dirname(os.path.abspath(__file__))


def buf(fig, dpi=220):
    b = BytesIO()
    fig.savefig(b, format='png', dpi=dpi, bbox_inches='tight',
                facecolor=fig.get_facecolor(), edgecolor='none')
    plt.close(fig)
    b.seek(0)
    return b


# ═══════════════════════════════════════
#  DIAGRAMS
# ═══════════════════════════════════════

def dia_hero():
    fig, ax = plt.subplots(figsize=(10, 2.2))
    fig.set_facecolor(BK)
    ax.set_facecolor(BK)
    ax.set_xlim(0,10); ax.set_ylim(0,2.2); ax.axis('off')

    ax.text(5, 1.65, 'PetStar', fontsize=30, fontweight='bold',
            color=WH, ha='center', fontfamily=FONT)
    ax.text(5, 1.18, '반려동물 사진 콘테스트 플랫폼  ·  백엔드 성능 최적화',
            fontsize=9.5, color=MD, ha='center', fontfamily=FONT)

    ms = [('p95 응답시간','97%↓','4.76s → 146ms'),
          ('처리량','28x','10.9 → 308 RPS'),
          ('에러율','0%','300 VUs 안정')]
    bw = 2.6; sx = 5-(3*bw+2*0.2)/2
    for i,(lb,v,dt) in enumerate(ms):
        x = sx+i*(bw+0.2)
        ax.add_patch(FancyBboxPatch((x,0.12), bw, 0.6, boxstyle="round,pad=0.05",
                                   facecolor=D1, edgecolor=D2, lw=0.8))
        ax.text(x+bw/2, 0.52, v, fontsize=15, fontweight='bold',
                color=WH, ha='center', fontfamily=FONT)
        ax.text(x+bw/2, 0.28, f'{lb}  {dt}', fontsize=6.5,
                color=MD, ha='center', fontfamily=FONT)
    return buf(fig)


def dia_waterfall():
    fig, ax = plt.subplots(figsize=(8, 3.3))
    fig.set_facecolor(WH)

    stages = ['Before', '1단계\nPageable', '2단계\nFetch Join', '3단계\n복합 인덱스']
    vals = [4760, 884, 234, 146]
    grays = [D1, D2, D3, BK]

    bars = ax.bar(stages, vals, color=grays, width=0.5, edgecolor=WH, lw=2)
    for bar,v in zip(bars,vals):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+100,
                f'{v:,}ms', ha='center', fontsize=10, fontweight='bold',
                color=BK, fontfamily=FONT)
    reds = ['','81%↓','73%↓','38%↓']
    for bar,r in zip(bars,reds):
        if r and bar.get_height()>300:
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()/2,
                    r, ha='center', fontsize=10, fontweight='bold',
                    color=WH, fontfamily=FONT)

    ax.annotate('', xy=(3,300), xytext=(0,4500),
                arrowprops=dict(arrowstyle='->', color=BK, lw=1.8,
                                connectionstyle='arc3,rad=-0.25'))
    ax.text(2.0, 3100, '97% 감소', fontsize=11, fontweight='bold',
            color=BK, ha='center', fontfamily=FONT, rotation=-28)

    ax.set_ylabel('p95 (ms)', fontsize=9, color=D2, fontfamily=FONT)
    ax.set_title('Ranking API 최적화 과정', fontsize=11, fontweight='bold',
                 color=BK, pad=10, fontfamily=FONT)
    for s in ['top','right']: ax.spines[s].set_visible(False)
    for s in ['left','bottom']: ax.spines[s].set_color(LN)
    ax.tick_params(colors=D3, labelsize=8)
    for l in ax.get_xticklabels()+ax.get_yticklabels(): l.set_fontfamily(FONT)
    fig.tight_layout()
    return buf(fig)


def dia_bars():
    fig, (a1,a2) = plt.subplots(1,2, figsize=(8, 2.6))
    fig.set_facecolor(WH)
    def sty(a):
        for s in ['top','right']: a.spines[s].set_visible(False)
        for s in ['left','bottom']: a.spines[s].set_color(LN)
        a.tick_params(colors=D3, labelsize=8)
        for l in a.get_xticklabels()+a.get_yticklabels(): l.set_fontfamily(FONT)

    cs = [D2, BK]
    b1 = a1.bar(['Before','After'], [21,1], color=cs, width=0.4, edgecolor=WH, lw=2)
    for bar,v in zip(b1,[21,1]):
        a1.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.5,
                f'{v}개', ha='center', fontsize=11, fontweight='bold',
                color=BK, fontfamily=FONT)
    a1.set_title('요청당 쿼리 수', fontsize=10, fontweight='bold', color=BK, fontfamily=FONT, pad=8)
    sty(a1)

    b2 = a2.bar(['Before','After'], [10.9,308], color=cs, width=0.4, edgecolor=WH, lw=2)
    for bar,v in zip(b2,[10.9,308]):
        a2.text(bar.get_x()+bar.get_width()/2, bar.get_height()+5,
                f'{v} RPS', ha='center', fontsize=11, fontweight='bold',
                color=BK, fontfamily=FONT)
    a2.set_title('처리량', fontsize=10, fontweight='bold', color=BK, fontfamily=FONT, pad=8)
    sty(a2)
    fig.tight_layout(w_pad=3)
    return buf(fig)


def dia_concurrency():
    fig, (a1,a2) = plt.subplots(1,2, figsize=(8, 3))
    fig.set_facecolor(WH)
    def sty(a):
        for s in ['top','right']: a.spines[s].set_visible(False)
        for s in ['left','bottom']: a.spines[s].set_color(LN)
        a.tick_params(colors=D3, labelsize=8)
        for l in a.get_xticklabels()+a.get_yticklabels(): l.set_fontfamily(FONT)

    lb = ['비관적 락\n(채택)','낙관적 락','Atomic\n+ Retry']
    cs = [BK, LT, LT]

    b1 = a1.bar(lb, [100,16,94], color=cs, width=0.45, edgecolor=WH, lw=2)
    for bar,v in zip(b1,[100,16,94]):
        a1.text(bar.get_x()+bar.get_width()/2, bar.get_height()+2,
                f'{v}%', ha='center', fontsize=11, fontweight='bold',
                color=BK, fontfamily=FONT)
    a1.set_ylim(0,115); a1.axhline(100, color=LN, ls='--', lw=0.8)
    a1.set_title('정합성 (50명 동시 투표)', fontsize=10, fontweight='bold', color=BK, fontfamily=FONT, pad=8)
    a1.set_ylabel('성공률 (%)', fontsize=8, color=D3, fontfamily=FONT)
    sty(a1)

    b2 = a2.bar(lb, [1730,885,2980], color=cs, width=0.45, edgecolor=WH, lw=2)
    for bar,v in zip(b2,[1730,885,2980]):
        a2.text(bar.get_x()+bar.get_width()/2, bar.get_height()+50,
                f'{v:,}ms', ha='center', fontsize=10, fontweight='bold',
                color=BK, fontfamily=FONT)
    a2.set_title('소요시간', fontsize=10, fontweight='bold', color=BK, fontfamily=FONT, pad=8)
    sty(a2)
    fig.tight_layout(w_pad=3)
    return buf(fig)


def _box(ax, x, y, w, h, label, sub='', filled=False):
    fc = BK if filled else WH
    ec = BK
    tc = WH if filled else BK
    sc = MD if filled else D3
    ax.add_patch(FancyBboxPatch((x,y), w, h, boxstyle="round,pad=0.05",
                                facecolor=fc, edgecolor=ec, lw=1.2))
    if sub:
        ax.text(x+w/2, y+h/2+0.12, label, ha='center', va='center',
                fontsize=8.5, fontweight='bold', color=tc, fontfamily=FONT)
        ax.text(x+w/2, y+h/2-0.14, sub, ha='center', va='center',
                fontsize=7, color=sc, fontfamily=FONT)
    else:
        ax.text(x+w/2, y+h/2, label, ha='center', va='center',
                fontsize=8.5, fontweight='bold', color=tc, fontfamily=FONT)

def _arr(ax, x1,y1, x2,y2, label='', dashed=False):
    ax.annotate('', xy=(x2,y2), xytext=(x1,y1),
                arrowprops=dict(arrowstyle='->', color=D2, lw=1.2,
                                linestyle='--' if dashed else '-'))
    if label:
        mx,my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx, my+0.16, label, ha='center', va='center', fontsize=7,
                color=D3, fontfamily=FONT,
                bbox=dict(boxstyle='round,pad=0.1', fc=WH, ec='none'))


def dia_retry():
    fig, (a1,a2) = plt.subplots(1,2, figsize=(9.5, 3.8))
    fig.set_facecolor(WH)

    a1.set_xlim(0,10); a1.set_ylim(0,10); a1.axis('off')
    a1.set_title('낙관적 락 — Retry Storm', fontsize=10, fontweight='bold',
                 color=BK, fontfamily=FONT, pad=6)

    for i,y in enumerate([8.2,6.6,5.0,3.4]):
        _box(a1, 0.2, y-0.2, 1.5, 0.5, f'User {i+1}')
        _arr(a1, 1.7, y, 4.2, y, '시도')
        a1.text(4.5, y, '×', fontsize=13, fontweight='bold', color=D2, ha='center')
        _arr(a1, 4.8, y, 6.8, y-0.15, '재시도', True)
        a1.text(7.1, y-0.15, '×', fontsize=13, fontweight='bold', color=D2, ha='center')
        _arr(a1, 7.4, y-0.15, 9.2, y-0.3, '', True)

    a1.add_patch(FancyBboxPatch((2.5,0.5), 5, 1.2, boxstyle="round,pad=0.1",
                                facecolor=BG, edgecolor=LN, lw=1))
    a1.text(5, 1.3, '충돌 → 재시도 → 재충돌 → 지수적 경쟁', fontsize=8.5,
            fontweight='bold', color=BK, ha='center', fontfamily=FONT)
    a1.text(5, 0.8, '성공률 16%  ·  885ms', fontsize=8,
            color=D3, ha='center', fontfamily=FONT)

    a2.set_xlim(0,10); a2.set_ylim(0,10); a2.axis('off')
    a2.set_title('비관적 락 — 순차 대기', fontsize=10, fontweight='bold',
                 color=BK, fontfamily=FONT, pad=6)

    a2.add_patch(FancyBboxPatch((0.5,7.0), 9, 1.0, boxstyle="round,pad=0.08",
                                facecolor=BG2, edgecolor=BK, lw=1.2))
    a2.text(5, 7.5, 'SELECT FOR UPDATE — 대기열', fontsize=9.5,
            fontweight='bold', color=BK, ha='center', fontfamily=FONT)

    for i,(lb,y) in enumerate(zip(
            ['User 1 → 처리 완료','User 2 → 처리 완료',
             'User 3 → 처리 완료','User N → 대기 중'],
            [5.5,4.2,2.9,1.6])):
        filled = i < 3
        fc = BK if filled else BG
        tc = WH if filled else D3
        ec = BK if filled else LN
        a2.add_patch(FancyBboxPatch((1.5,y-0.2), 7, 0.5, boxstyle="round,pad=0.04",
                                   facecolor=fc, edgecolor=ec, lw=1))
        a2.text(5, y+0.05, lb, ha='center', va='center', fontsize=8.5,
                fontweight='bold', color=tc, fontfamily=FONT)
        if i < 3: _arr(a2, 5, y-0.25, 5, y-0.55)

    a2.add_patch(FancyBboxPatch((2.5,0.3), 5, 0.6, boxstyle="round,pad=0.08",
                                facecolor=BG, edgecolor=LN, lw=1))
    a2.text(5, 0.6, '성공률 100%  ·  충돌 없음', fontsize=9,
            fontweight='bold', color=BK, ha='center', fontfamily=FONT)

    fig.tight_layout(w_pad=2)
    return buf(fig)


def dia_arch():
    fig, ax = plt.subplots(figsize=(10, 4.8))
    fig.set_facecolor(WH)
    ax.set_xlim(0,12); ax.set_ylim(0,6.2); ax.axis('off')

    ax.text(6, 5.95, '비동기 투표 시스템 아키텍처', fontsize=13, fontweight='bold',
            color=BK, ha='center', fontfamily=FONT)

    # Fast path
    ax.add_patch(FancyBboxPatch((0.2,3.0), 11.6, 2.5, boxstyle="round,pad=0.12",
                                facecolor=BG, edgecolor=LN, lw=1, linestyle='--'))
    ax.text(1.1, 5.2, 'Fast Path', fontsize=8.5, fontweight='bold', color=D2, fontfamily=FONT)

    _box(ax, 0.5, 3.4, 1.5, 1.0, '클라이언트', '투표 요청', True)
    _box(ax, 3.0, 3.4, 2.0, 1.0, 'Spring Boot', 'API Server')
    _box(ax, 6.2, 3.4, 2.0, 1.0, 'Redis', 'Sorted Set')

    _arr(ax, 2.0, 3.95, 3.0, 3.95, '투표')
    _arr(ax, 5.0, 4.05, 6.2, 4.05, 'ZINCRBY')
    _arr(ax, 6.2, 3.7, 5.0, 3.7, '반영')
    _arr(ax, 3.0, 3.7, 2.0, 3.7, '응답')

    ax.add_patch(FancyBboxPatch((9.0, 3.6), 2.4, 0.6, boxstyle="round,pad=0.06",
                                facecolor=BK, edgecolor=BK, lw=1))
    ax.text(10.2, 3.9, '~5ms 응답', fontsize=10, fontweight='bold',
            color=WH, ha='center', fontfamily=FONT)

    # Slow path
    ax.add_patch(FancyBboxPatch((0.2,0.15), 11.6, 2.4, boxstyle="round,pad=0.12",
                                facecolor=WH, edgecolor=LN, lw=1, linestyle='--'))
    ax.text(1.1, 2.25, 'Slow Path', fontsize=8.5, fontweight='bold', color=D3, fontfamily=FONT)

    _box(ax, 0.5, 0.5, 2.0, 0.9, 'SQS Queue', '메시지')
    _box(ax, 3.5, 0.5, 2.0, 0.9, 'Consumer', '비동기 처리')
    _box(ax, 6.5, 0.5, 2.0, 0.9, 'MySQL RDS', 'Source of Truth', True)

    _arr(ax, 4.0, 3.4, 1.5, 1.4, 'Fire & Forget', True)
    _arr(ax, 2.5, 0.95, 3.5, 0.95, '폴링')
    _arr(ax, 5.5, 0.95, 6.5, 0.95, 'DB 저장')

    _box(ax, 9.2, 0.7, 2.3, 0.6, 'Scheduler 5분', '정합성 검증')
    _arr(ax, 9.2, 1.1, 8.5, 0.95, '', True)
    _arr(ax, 10.0, 1.3, 7.8, 3.4, '', True)

    fig.tight_layout()
    return buf(fig)


def dia_reversal():
    fig, ax = plt.subplots(figsize=(8, 3.5))
    fig.set_facecolor(WH)

    ph = ['Phase 6\n비관적 락','Phase 10\nSQS 비동기',
          'Phase 10-2\nRedis 중복체크','Phase 10-3\nPooling']
    vs = [1730, 516, 1300, 658]
    cs = [LT, D3, BK, D2]

    bars = ax.bar(ph, vs, color=cs, width=0.5, edgecolor=WH, lw=2)
    for bar,v in zip(bars,vs):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+30,
                f'{v:,}ms', ha='center', fontsize=10, fontweight='bold',
                color=BK, fontfamily=FONT)

    ax.annotate('성능 역전', xy=(2,1300), xytext=(2.7,1550),
                fontsize=11, fontweight='bold', color=BK, fontfamily=FONT,
                arrowprops=dict(arrowstyle='->', color=BK, lw=1.5))
    ax.add_patch(plt.Rectangle((1.55,0), 0.9, 1350, fill=False,
                               edgecolor=BK, lw=1.5, ls='--'))

    ax.set_title('투표 API p95 변화 추이', fontsize=11, fontweight='bold',
                 color=BK, pad=10, fontfamily=FONT)
    ax.set_ylabel('p95 (ms)', fontsize=9, color=D3, fontfamily=FONT)
    for s in ['top','right']: ax.spines[s].set_visible(False)
    for s in ['left','bottom']: ax.spines[s].set_color(LN)
    ax.tick_params(colors=D3, labelsize=8)
    for l in ax.get_xticklabels()+ax.get_yticklabels(): l.set_fontfamily(FONT)
    fig.tight_layout()
    return buf(fig)


def dia_conn():
    fig, (a1,a2) = plt.subplots(1,2, figsize=(10, 4))
    fig.set_facecolor(WH)

    # Lettuce
    a1.set_xlim(0,10); a1.set_ylim(0,10); a1.axis('off')
    a1.set_title('Lettuce — 단일 커넥션', fontsize=10, fontweight='bold',
                 color=BK, fontfamily=FONT, pad=6)

    for i in range(5):
        y = 8.5 - i*1.2
        _box(a1, 0.2, y-0.2, 1.5, 0.5, f'요청 {i+1}')
        a1.annotate('', xy=(4.3,5.05), xytext=(1.7,y),
                     arrowprops=dict(arrowstyle='->', color=LN, lw=0.7))

    a1.add_patch(FancyBboxPatch((4.2,4.0), 2.2, 1.8, boxstyle="round,pad=0.08",
                                facecolor=BK, edgecolor=BK, lw=1.5))
    a1.text(5.3, 5.15, 'TCP 1개', fontsize=11, fontweight='bold',
            color=WH, ha='center', fontfamily=FONT)
    a1.text(5.3, 4.6, '직렬화 병목', fontsize=8, color=MD, ha='center', fontfamily=FONT)

    _box(a1, 7.2, 4.3, 2.0, 1.2, 'Redis', '', True)
    _arr(a1, 6.4, 5.0, 7.2, 5.0)

    a1.add_patch(FancyBboxPatch((2.8,1.0), 4.4, 0.8, boxstyle="round,pad=0.08",
                                facecolor=BG, edgecolor=LN, lw=1))
    a1.text(5, 1.4, 'p95: 1,300ms', fontsize=12, fontweight='bold',
            color=BK, ha='center', fontfamily=FONT)

    # HikariCP
    a2.set_xlim(0,10); a2.set_ylim(0,10); a2.axis('off')
    a2.set_title('HikariCP — 30개 커넥션 풀', fontsize=10, fontweight='bold',
                 color=BK, fontfamily=FONT, pad=6)

    for i in range(5):
        y = 8.5 - i*1.2
        _box(a2, 0.2, y-0.2, 1.5, 0.5, f'요청 {i+1}')

    conn_ys = [7.8, 6.2, 4.6, 3.0]
    a2.add_patch(FancyBboxPatch((3.3,2.2), 2.4, 6.5, boxstyle="round,pad=0.12",
                                facecolor='none', edgecolor=BK, lw=1.3, ls='--'))
    a2.text(4.5, 2.45, '30개 풀', fontsize=8, fontweight='bold',
            color=D2, ha='center', fontfamily=FONT)

    for i,cy in enumerate(conn_ys):
        a2.add_patch(FancyBboxPatch((3.6,cy-0.2), 1.8, 0.5, boxstyle="round,pad=0.04",
                                   facecolor=BK, edgecolor=BK, lw=1))
        a2.text(4.5, cy+0.05, f'Conn {i+1}', fontsize=7.5, fontweight='bold',
                color=WH, ha='center', fontfamily=FONT)
        a2.annotate('', xy=(7.0,5.05), xytext=(5.4,cy+0.05),
                     arrowprops=dict(arrowstyle='->', color=LN, lw=0.7))

    for i in range(5):
        y = 8.5-i*1.2; t = conn_ys[min(i,3)]
        a2.annotate('', xy=(3.6,t+0.05), xytext=(1.7,y),
                     arrowprops=dict(arrowstyle='->', color=LN, lw=0.7))

    _box(a2, 7.0, 4.2, 2.2, 1.6, 'MySQL', '병렬 처리', True)

    a2.add_patch(FancyBboxPatch((2.8,1.0), 4.4, 0.8, boxstyle="round,pad=0.08",
                                facecolor=BG, edgecolor=LN, lw=1))
    a2.text(5, 1.4, 'p95: 516ms', fontsize=12, fontweight='bold',
            color=BK, ha='center', fontfamily=FONT)

    fig.tight_layout(w_pad=2)
    return buf(fig)


# ═══════════════════════════════════════
#  DOCX
# ═══════════════════════════════════════

def hx(h):
    h = h.lstrip('#')
    return RGBColor(int(h[:2],16), int(h[2:4],16), int(h[4:6],16))

def _shd(cell, c):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{c.lstrip("#")}"/>'))

def _r(p, t, sz=9.5, b=False, c=BK, it=False):
    r = p.add_run(t)
    r.font.size=Pt(sz); r.font.bold=b; r.font.italic=it
    r.font.color.rgb=hx(c); r.font.name=FONT
    r._element.get_or_add_rPr().append(
        parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT}"/>'))
    return r

def pa(doc, t='', sz=9.5, b=False, c=BK, al=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=4):
    p = doc.add_paragraph(); p.alignment=al
    p.paragraph_format.space_before=Pt(sb)
    p.paragraph_format.space_after=Pt(sa)
    if t: _r(p, t, sz, b, c)
    return p

def sec_head(doc, num, title, sub):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(2)
    rn = _r(p, f' {num} ', 12, True, WH)
    rn._element.get_or_add_rPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{BK.lstrip("#")}"/>'))
    _r(p, '  ', 12)
    _r(p, title, 14, True, BK)
    pa(doc, sub, 9.5, False, D3, sb=0, sa=10)

def lbl(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    _r(p, f'— {t}', 9, True, D2)

def ins(doc, t):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(8)
    p.paragraph_format.left_indent = Cm(0.4)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="16" w:space="8" w:color="{D2.lstrip("#")}"/>'
        f'</w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{BG.lstrip("#")}"/>'))
    _r(p, t, 9, False, BK)

def tbl(doc, hds, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(hds))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = t._tbl.tblPr if t._tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{LN.lstrip("#")}"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{LN.lstrip("#")}"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{LN.lstrip("#")}"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'))
    for i,h in enumerate(hds):
        c = t.rows[0].cells[i]; c.text=''
        _shd(c, BK)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _r(p, h, 8.5, True, WH)
    for ri,row in enumerate(rows):
        for ci,v in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text=''
            if ri%2==1: _shd(c, BG)
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _r(p, str(v), 8.5, False, BK)

def div(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before=Pt(4); p.paragraph_format.space_after=Pt(4)
    p._element.get_or_add_pPr().append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{LN.lstrip("#")}"/>'
        f'</w:pBdr>'))


def build():
    doc = Document()
    s = doc.sections[0]
    s.page_width=Cm(21); s.page_height=Cm(29.7)
    s.top_margin=Cm(1.5); s.bottom_margin=Cm(1.5)
    s.left_margin=Cm(2.0); s.right_margin=Cm(2.0)
    st = doc.styles['Normal']
    st.font.name=FONT; st.font.size=Pt(9.5); st.font.color.rgb=hx(BK)
    st.element.get_or_add_rPr().append(
        parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT}"/>'))

    # HERO
    doc.add_picture(dia_hero(), width=Cm(17))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = pa(doc, '', sa=1, sb=4, al=WD_ALIGN_PARAGRAPH.CENTER)
    for i,t in enumerate(['Java 21','Spring Boot 3.5','MySQL 8.0','Redis 7','AWS SQS',
                          'QueryDSL','Docker','Terraform','k6','Prometheus']):
        _r(p, t, 8, False, D2)
        if i<9: _r(p, '  ·  ', 8, False, LN)

    pa(doc, '반려동물 사진 콘테스트 플랫폼 — 챌린지별 투표 → 실시간 랭킹',
       9, False, D3, WD_ALIGN_PARAGRAPH.CENTER, 2, 2)
    pa(doc, '10K 회원 · 100K 엔트리 · 1M 투표  |  EC2 t3.small · RDS db.t3.micro',
       8, False, MD, WD_ALIGN_PARAGRAPH.CENTER, 0, 6)
    div(doc)

    # ══ 01 ══
    sec_head(doc, '01', 'DB 쿼리 최적화', 'Ranking API p95 4.76s → 146ms (97% 감소)')

    lbl(doc, 'PROBLEM')
    pa(doc, 'k6 부하 테스트(50 VUs)에서 랭킹 API p95가 4.76초로 측정. '
       'Grafana + p6spy 분석 결과 세 가지 병목 확인.', sb=2)
    tbl(doc, ['병목','원인','영향'],
        [['데이터 전송량','Java Stream.limit() — DB에서 전체 전송','네트워크/메모리 낭비'],
         ['N+1 쿼리','연관 엔티티 개별 SELECT — 요청당 21개','커넥션 점유 증가'],
         ['filesort','vote_count 인덱스 없음','CPU 사용률 증가']])

    lbl(doc, 'SOLUTION')
    pa(doc, '상위 레이어부터 순차 최적화 — 상위 병목을 먼저 제거해야 하위 레이어 효과를 정확히 측정 가능.',
       9.5, True, sb=2)
    doc.add_picture(dia_waterfall(), width=Cm(14))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl(doc, ['단계','최적화','Before','After','개선율'],
        [['1단계','Pageable — DB LIMIT','p95 4,760ms','p95 884ms','81%↓'],
         ['2단계','Fetch Join — N+1 해결 (21→1)','p95 884ms','p95 234ms','73%↓'],
         ['3단계','복합 인덱스 (challenge_id, vote_count DESC)','p95 234ms','p95 146ms','38%↓']])
    doc.add_picture(dia_bars(), width=Cm(14))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    lbl(doc, 'RESULT')
    tbl(doc, ['지표','Before (50 VUs)','After (300 VUs)','개선'],
        [['랭킹 API p95','4,760ms','178ms','97%↓'],
         ['쿼리 수','21개/요청','1개/요청','95%↓'],
         ['처리량','10.9 RPS','308 RPS','28x'],
         ['에러율','-','0%','300 VUs 안정']])
    ins(doc, '최적화 순서가 중요하다. 1단계에서 데이터 전송량을 줄이지 않으면 '
        '2단계 Fetch Join 효과가 과대/과소 측정된다. 측정 가능한 최적화를 위해 '
        '상위 레이어부터 제거하는 원칙을 적용했다.')
    div(doc)

    # ══ 02 ══
    sec_head(doc, '02', '동시성 제어', '3가지 전략 비교 → 비관적 락 채택')

    lbl(doc, 'PROBLEM')
    pa(doc, '투표 기능에서 두 가지 레이스 컨디션 발생.', sb=2)
    tbl(doc, ['레이스 컨디션','원인','결과'],
        [['Check-then-Act','SELECT ↔ INSERT 갭','중복 투표'],
         ['Lost Update','동시 voteCount++ 유실','투표 수 불일치']])
    pa(doc, 'synchronized: JVM 레벨, 다중 인스턴스 무효. '
       '@Transactional 단독: REPEATABLE READ에서 Lost Update 미방지.', 9, False, D3, sb=4)

    lbl(doc, 'SOLUTION')
    pa(doc, '동일 조건(50명 동시 투표, 동일 Entry)에서 3가지 전략 구현 및 테스트.', sb=2)
    doc.add_picture(dia_concurrency(), width=Cm(14))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl(doc, ['전략','성공률','소요시간','문제점'],
        [['비관적 락 (채택)','100%','1,730ms','없음'],
         ['낙관적 락 (@Version)','16%','885ms','42/50 실패'],
         ['Atomic + Retry','94%','2,980ms','Retry Storm']])
    doc.add_picture(dia_retry(), width=Cm(15))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    lbl(doc, 'TROUBLESHOOTING')
    tbl(doc, ['문제','원인','해결'],
        [['@Version NULL','기존 데이터 version 없음','ALTER TABLE 기본값 마이그레이션'],
         ['@Transactional 미적용','내부 호출 시 AOP 프록시 우회','@Lazy 자기 주입']])
    ins(doc, '"느린 비관적 락이 역설적으로 빠르다." 고경쟁 환경에서 낙관적 접근은 '
        '충돌 → 재시도 → 재충돌의 Retry Storm을 일으켜 전체 응답시간이 증가한다.')
    div(doc)

    # ══ 03 ══
    sec_head(doc, '03', '비동기 투표 아키텍처', '"빠른 기술 ≠ 빠른 시스템"')

    lbl(doc, 'PROBLEM')
    pa(doc, '비관적 락 p95 1.73초 — 동기 DB 쓰기 병목.', sb=2)

    lbl(doc, 'SOLUTION — 읽기/쓰기 분리')
    doc.add_picture(dia_arch(), width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    pa(doc, 'Fast Path: Redis ZINCRBY → 즉시 응답 (~5ms)\n'
       'Slow Path: SQS → Consumer → DB 영속화\n'
       '정합성: 5분 Scheduler Redis ↔ DB 검증', sb=4)

    lbl(doc, 'MQ 선택')
    tbl(doc, ['기준','Kafka (MSK)','RabbitMQ','SQS (채택)'],
        [['월 비용','$200+','$50-100','$1'],
         ['운영','ZooKeeper','Erlang','Zero-Ops'],
         ['확장','수동','수동','자동']])
    pa(doc, '순서 보장 불필요 + 멱등성 설계(중복 체크 + UK) → SQS Standard. '
       'DLQ(3회 재시도, 14일 보관).', 9, False, D3, sb=4)

    lbl(doc, 'DISCOVERY — 성능 역전')
    pa(doc, '중복 체크를 DB → Redis(SISMEMBER)로 이관 → 2.5배 느려짐.', 9.5, True, sb=4)
    doc.add_picture(dia_reversal(), width=Cm(14))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    lbl(doc, 'ROOT CAUSE')
    doc.add_picture(dia_conn(), width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl(doc, ['','Lettuce (Redis)','HikariCP (MySQL)'],
        [['커넥션','단일 TCP 멀티플렉싱','30개 커넥션 풀'],
         ['50 동시 요청','직렬화 → 병목','병렬 분산'],
         ['결과','p95 1,300ms','p95 516ms']])

    lbl(doc, 'RESOLUTION')
    pa(doc, 'Lettuce Connection Pooling(commons-pool2) + Redis Pipeline 적용. '
       'DB 중복 체크를 유지하는 Hybrid 전략 채택.', sb=4)

    lbl(doc, 'RESULT')
    tbl(doc, ['단계','p95','설명'],
        [['비관적 락','1,730ms','동기 DB 병목'],
         ['SQS 비동기','516ms','70%↓'],
         ['Redis 중복체크','1,300ms','성능 역전'],
         ['Pooling+Pipeline','658ms','최종 62%↓']])
    ins(doc, '"빠른 기술 ≠ 빠른 시스템." Redis O(1)보다 HikariCP 30개 커넥션의 병렬 처리가 '
        '50 동시 사용자에서 더 효과적이었다. 커넥션 모델이 성능을 결정한다.')
    div(doc)

    # Summary
    pa(doc, '기술적 의사결정 요약', 13, True, BK, sb=16, sa=8)
    tbl(doc, ['영역','선택','근거'],
        [['캐시','Redis (Global)','Scale-out 일관성'],
         ['동시성','비관적 락','Hot-spot 100% 정합성'],
         ['MQ','SQS Standard','$1/mo, Zero-Ops, 멱등성 설계'],
         ['중복 체크','DB (Hybrid)','커넥션 풀 병렬이 단일 커넥션보다 유리'],
         ['정합성','DB = Source of Truth','Redis 장애 시 DB 복구 가능']])

    out = os.path.join(OUT, 'PetStar_Portfolio_v3.docx')
    doc.save(out)
    print(f'Saved: {out}')


if __name__ == '__main__':
    build()
