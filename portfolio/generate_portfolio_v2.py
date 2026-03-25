#!/usr/bin/env python3
"""PetStar Portfolio DOCX Generator v2 — Minimal, Engineer-tone"""

import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch
import numpy as np
from io import BytesIO
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Palette: Navy + Blue accent + Grays only ──
INK     = '#1E293B'   # primary text
DIM     = '#64748B'   # secondary text
MUTE    = '#94A3B8'   # tertiary
LINE    = '#CBD5E1'   # borders
BG      = '#F8FAFC'   # subtle bg
WHITE   = '#FFFFFF'
ACCENT  = '#2563EB'   # the ONE accent color
ACC_L   = '#DBEAFE'   # accent light (backgrounds only)
NAVY    = '#0F172A'   # darkest

FONT = 'Apple SD Gothic Neo'
plt.rcParams['font.family'] = FONT
plt.rcParams['axes.unicode_minus'] = False

OUT_DIR = os.path.dirname(os.path.abspath(__file__))


def fig_to_bytes(fig, dpi=220):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                facecolor=fig.get_facecolor(), edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


# ═══════════════════════════════════════════
#  DIAGRAMS — monochrome engineering style
# ═══════════════════════════════════════════

def create_hero():
    fig, ax = plt.subplots(figsize=(10, 2.4))
    fig.set_facecolor(NAVY)
    ax.set_facecolor(NAVY)
    ax.set_xlim(0, 10); ax.set_ylim(0, 2.4)
    ax.axis('off')

    ax.text(5, 1.75, 'PetStar', fontsize=32, fontweight='bold',
            color=WHITE, ha='center', fontfamily=FONT)
    ax.text(5, 1.28, '반려동물 사진 콘테스트 플랫폼  ·  백엔드 성능 최적화',
            fontsize=10, color=MUTE, ha='center', fontfamily=FONT)

    metrics = [
        ('p95 응답시간', '97%↓', '4.76s → 146ms'),
        ('처리량', '28x', '10.9 → 308 RPS'),
        ('에러율', '0%', '300 VUs 안정'),
    ]
    bw, bh = 2.6, 0.65
    sx = 5 - (3*bw + 2*0.25)/2
    for i, (label, val, detail) in enumerate(metrics):
        x = sx + i*(bw+0.25)
        ax.add_patch(FancyBboxPatch((x, 0.15), bw, bh, boxstyle="round,pad=0.06",
                                   facecolor='#1E3A5F', edgecolor='#334155', lw=0.8))
        ax.text(x+bw/2, 0.58, val, fontsize=16, fontweight='bold',
                color=WHITE, ha='center', fontfamily=FONT)
        ax.text(x+bw/2, 0.32, f'{label}  {detail}', fontsize=7,
                color=MUTE, ha='center', fontfamily=FONT)

    return fig_to_bytes(fig)


def create_waterfall():
    fig, ax = plt.subplots(figsize=(8.5, 3.5))
    fig.set_facecolor(WHITE)

    stages = ['Before', '1단계\nPageable', '2단계\nFetch Join', '3단계\n복합 인덱스']
    values = [4760, 884, 234, 146]
    # Monochrome gradient: darkest for worst, lightest for best
    grays = ['#334155', '#475569', '#64748B', ACCENT]

    bars = ax.bar(stages, values, color=grays, width=0.5, edgecolor=WHITE, linewidth=2)
    for bar, val in zip(bars, values):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+100,
                f'{val:,}ms', ha='center', fontsize=11, fontweight='bold',
                color=INK, fontfamily=FONT)

    # Reduction labels inside bars
    reds = ['', '81%↓', '73%↓', '38%↓']
    for bar, r in zip(bars, reds):
        if r and bar.get_height() > 300:
            ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()/2,
                    r, ha='center', fontsize=11, fontweight='bold',
                    color=WHITE, fontfamily=FONT)

    ax.annotate('', xy=(3, 300), xytext=(0, 4500),
                arrowprops=dict(arrowstyle='->', color=ACCENT, lw=2,
                                connectionstyle='arc3,rad=-0.25'))
    ax.text(2.0, 3100, '97% 감소', fontsize=12, fontweight='bold',
            color=ACCENT, ha='center', fontfamily=FONT, rotation=-28)

    ax.set_ylabel('p95 응답시간 (ms)', fontsize=9, color=DIM, fontfamily=FONT)
    ax.set_title('Ranking API 최적화 과정', fontsize=12, fontweight='bold',
                 color=INK, pad=12, fontfamily=FONT)
    for spine in ['top','right']:
        ax.spines[spine].set_visible(False)
    for spine in ['left','bottom']:
        ax.spines[spine].set_color(LINE)
    ax.tick_params(colors=DIM, labelsize=8)
    for l in ax.get_xticklabels()+ax.get_yticklabels():
        l.set_fontfamily(FONT)
    fig.tight_layout()
    return fig_to_bytes(fig)


def create_query_throughput():
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(8.5, 2.8))
    fig.set_facecolor(WHITE)

    def style_ax(ax):
        for s in ['top','right']: ax.spines[s].set_visible(False)
        for s in ['left','bottom']: ax.spines[s].set_color(LINE)
        ax.tick_params(colors=DIM, labelsize=8)
        for l in ax.get_xticklabels()+ax.get_yticklabels(): l.set_fontfamily(FONT)

    # Queries
    cats = ['Before', 'After']
    vals = [21, 1]
    cols = ['#475569', ACCENT]
    b1 = ax1.bar(cats, vals, color=cols, width=0.4, edgecolor=WHITE, lw=2)
    for bar,v in zip(b1,vals):
        ax1.text(bar.get_x()+bar.get_width()/2, bar.get_height()+0.5,
                 f'{v}개', ha='center', fontsize=12, fontweight='bold',
                 color=INK, fontfamily=FONT)
    ax1.set_title('요청당 쿼리 수', fontsize=11, fontweight='bold',
                  color=INK, fontfamily=FONT, pad=8)
    style_ax(ax1)

    # Throughput
    b2 = ax2.bar(cats, [10.9, 308], color=cols, width=0.4, edgecolor=WHITE, lw=2)
    for bar,v in zip(b2,[10.9,308]):
        ax2.text(bar.get_x()+bar.get_width()/2, bar.get_height()+5,
                 f'{v} RPS', ha='center', fontsize=12, fontweight='bold',
                 color=INK, fontfamily=FONT)
    ax2.set_title('처리량 (Requests/sec)', fontsize=11, fontweight='bold',
                  color=INK, fontfamily=FONT, pad=8)
    style_ax(ax2)

    fig.tight_layout(w_pad=3)
    return fig_to_bytes(fig)


def create_concurrency():
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(8.5, 3.2))
    fig.set_facecolor(WHITE)

    def style_ax(ax):
        for s in ['top','right']: ax.spines[s].set_visible(False)
        for s in ['left','bottom']: ax.spines[s].set_color(LINE)
        ax.tick_params(colors=DIM, labelsize=8)
        for l in ax.get_xticklabels()+ax.get_yticklabels(): l.set_fontfamily(FONT)

    labels = ['비관적 락\n(채택)', '낙관적 락', 'Atomic\n+ Retry']
    cols = [ACCENT, '#94A3B8', '#94A3B8']

    # Success rate
    rates = [100, 16, 94]
    b1 = ax1.bar(labels, rates, color=cols, width=0.45, edgecolor=WHITE, lw=2)
    for bar,v in zip(b1, rates):
        ax1.text(bar.get_x()+bar.get_width()/2, bar.get_height()+2,
                 f'{v}%', ha='center', fontsize=12, fontweight='bold',
                 color=INK, fontfamily=FONT)
    ax1.set_ylim(0, 115)
    ax1.axhline(100, color=LINE, ls='--', lw=0.8)
    ax1.set_title('정합성 (50명 동시 투표)', fontsize=11, fontweight='bold',
                  color=INK, fontfamily=FONT, pad=8)
    ax1.set_ylabel('성공률 (%)', fontsize=9, color=DIM, fontfamily=FONT)
    style_ax(ax1)

    # Time
    times = [1730, 885, 2980]
    b2 = ax2.bar(labels, times, color=cols, width=0.45, edgecolor=WHITE, lw=2)
    for bar,v in zip(b2, times):
        ax2.text(bar.get_x()+bar.get_width()/2, bar.get_height()+50,
                 f'{v:,}ms', ha='center', fontsize=11, fontweight='bold',
                 color=INK, fontfamily=FONT)
    ax2.set_title('소요시간', fontsize=11, fontweight='bold',
                  color=INK, fontfamily=FONT, pad=8)
    ax2.set_ylabel('ms', fontsize=9, color=DIM, fontfamily=FONT)
    style_ax(ax2)

    fig.tight_layout(w_pad=3)
    return fig_to_bytes(fig)


def _box(ax, x, y, w, h, label, sub='', filled=False):
    """Draw a clean engineering-style box"""
    fc = NAVY if filled else WHITE
    ec = NAVY if filled else '#475569'
    tc = WHITE if filled else INK
    sc = MUTE if filled else DIM
    ax.add_patch(FancyBboxPatch((x,y), w, h, boxstyle="round,pad=0.06",
                                facecolor=fc, edgecolor=ec, lw=1.2))
    if sub:
        ax.text(x+w/2, y+h/2+0.12, label, ha='center', va='center',
                fontsize=9, fontweight='bold', color=tc, fontfamily=FONT)
        ax.text(x+w/2, y+h/2-0.14, sub, ha='center', va='center',
                fontsize=7, color=sc, fontfamily=FONT)
    else:
        ax.text(x+w/2, y+h/2, label, ha='center', va='center',
                fontsize=9, fontweight='bold', color=tc, fontfamily=FONT)


def _arrow(ax, x1, y1, x2, y2, label='', dashed=False):
    ax.annotate('', xy=(x2,y2), xytext=(x1,y1),
                arrowprops=dict(arrowstyle='->', color='#475569', lw=1.3,
                                linestyle='--' if dashed else '-'))
    if label:
        mx, my = (x1+x2)/2, (y1+y2)/2
        ax.text(mx, my+0.16, label, ha='center', va='center', fontsize=7,
                color=DIM, fontfamily=FONT,
                bbox=dict(boxstyle='round,pad=0.1', fc=WHITE, ec='none'))


def create_retry_vs_pessimistic():
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9.5, 4))
    fig.set_facecolor(WHITE)

    # ── Left: Retry Storm ──
    ax1.set_xlim(0,10); ax1.set_ylim(0,10); ax1.axis('off')
    ax1.set_title('낙관적 락 — Retry Storm', fontsize=11, fontweight='bold',
                  color=INK, fontfamily=FONT, pad=8)

    for i, y in enumerate([8.2, 6.6, 5.0, 3.4]):
        _box(ax1, 0.2, y-0.2, 1.5, 0.5, f'User {i+1}')
        # Attempt 1
        _arrow(ax1, 1.7, y, 4.2, y, '시도')
        ax1.text(4.5, y, '×', fontsize=14, fontweight='bold',
                 color='#475569', ha='center', va='center', fontfamily=FONT)
        # Retry
        _arrow(ax1, 4.8, y, 6.8, y-0.15, '재시도', dashed=True)
        ax1.text(7.1, y-0.15, '×', fontsize=14, fontweight='bold',
                 color='#475569', ha='center', va='center', fontfamily=FONT)
        # Re-retry
        _arrow(ax1, 7.4, y-0.15, 9.2, y-0.3, '', dashed=True)
        ax1.text(9.5, y-0.3, '×', fontsize=12, color='#475569',
                 ha='center', va='center', fontfamily=FONT)

    ax1.add_patch(FancyBboxPatch((2.5, 0.5), 5, 1.2, boxstyle="round,pad=0.1",
                                facecolor=BG, edgecolor=LINE, lw=1))
    ax1.text(5, 1.3, '충돌 → 재시도 → 재충돌 → 지수적 경쟁', fontsize=9,
             fontweight='bold', color=INK, ha='center', fontfamily=FONT)
    ax1.text(5, 0.8, '성공률 16%  ·  소요시간 885ms', fontsize=8,
             color=DIM, ha='center', fontfamily=FONT)

    # ── Right: Pessimistic Lock ──
    ax2.set_xlim(0,10); ax2.set_ylim(0,10); ax2.axis('off')
    ax2.set_title('비관적 락 — 순차 대기', fontsize=11, fontweight='bold',
                  color=INK, fontfamily=FONT, pad=8)

    # Queue
    ax2.add_patch(FancyBboxPatch((0.5, 7.0), 9, 1.2, boxstyle="round,pad=0.1",
                                facecolor=ACC_L, edgecolor=ACCENT, lw=1.2))
    ax2.text(5, 7.6, 'SELECT FOR UPDATE — 대기열', fontsize=10,
             fontweight='bold', color=ACCENT, ha='center', fontfamily=FONT)

    for i, (label, y) in enumerate(zip(
            ['User 1 → 처리 완료', 'User 2 → 처리 완료',
             'User 3 → 처리 완료', 'User N → 대기 중'],
            [5.5, 4.2, 2.9, 1.6])):
        filled = (i < 3)
        fc = ACCENT if filled else BG
        tc = WHITE if filled else DIM
        ec = ACCENT if filled else LINE
        ax2.add_patch(FancyBboxPatch((1.5, y-0.2), 7, 0.55, boxstyle="round,pad=0.05",
                                    facecolor=fc, edgecolor=ec, lw=1))
        ax2.text(5, y+0.07, label, ha='center', va='center', fontsize=9,
                 fontweight='bold', color=tc, fontfamily=FONT)
        if i < 3:
            _arrow(ax2, 5, y-0.25, 5, y-0.55)

    ax2.add_patch(FancyBboxPatch((2.5, 0.15), 5, 0.7, boxstyle="round,pad=0.1",
                                facecolor=BG, edgecolor=LINE, lw=1))
    ax2.text(5, 0.5, '성공률 100%  ·  충돌 없음  ·  예측 가능', fontsize=9,
             fontweight='bold', color=ACCENT, ha='center', fontfamily=FONT)

    fig.tight_layout(w_pad=2)
    return fig_to_bytes(fig)


def create_architecture():
    fig, ax = plt.subplots(figsize=(10, 5))
    fig.set_facecolor(WHITE)
    ax.set_xlim(0,12); ax.set_ylim(0,6.5); ax.axis('off')

    ax.text(6, 6.2, '비동기 투표 시스템 아키텍처', fontsize=14, fontweight='bold',
            color=INK, ha='center', fontfamily=FONT)

    # ── Fast path region ──
    ax.add_patch(FancyBboxPatch((0.2, 3.2), 11.6, 2.6, boxstyle="round,pad=0.12",
                                facecolor=ACC_L, edgecolor=ACCENT, lw=1, alpha=0.3))
    ax.text(1.2, 5.5, 'Fast Path — 즉시 응답', fontsize=9, fontweight='bold',
            color=ACCENT, fontfamily=FONT)

    _box(ax, 0.5, 3.6, 1.6, 1.1, '클라이언트', '투표 요청', filled=True)
    _box(ax, 3.0, 3.6, 2.0, 1.1, 'Spring Boot', 'API Server')
    _box(ax, 6.2, 3.6, 2.0, 1.1, 'Redis', 'Sorted Set')

    _arrow(ax, 2.1, 4.2, 3.0, 4.2, '투표')
    _arrow(ax, 5.0, 4.3, 6.2, 4.3, 'ZINCRBY')
    _arrow(ax, 6.2, 3.9, 5.0, 3.9, '랭킹 반영')
    _arrow(ax, 3.0, 3.9, 2.1, 3.9, '응답')

    # Response badge
    ax.add_patch(FancyBboxPatch((9.0, 3.8), 2.5, 0.7, boxstyle="round,pad=0.08",
                                facecolor=ACCENT, edgecolor=ACCENT, lw=1))
    ax.text(10.25, 4.15, '~5ms 응답', fontsize=11, fontweight='bold',
            color=WHITE, ha='center', fontfamily=FONT)

    # ── Slow path region ──
    ax.add_patch(FancyBboxPatch((0.2, 0.2), 11.6, 2.5, boxstyle="round,pad=0.12",
                                facecolor=BG, edgecolor=LINE, lw=1))
    ax.text(1.2, 2.4, 'Slow Path — 비동기 영속화', fontsize=9, fontweight='bold',
            color=DIM, fontfamily=FONT)

    _box(ax, 0.5, 0.6, 2.0, 1.0, 'SQS Queue', '메시지')
    _box(ax, 3.5, 0.6, 2.0, 1.0, 'Consumer', '비동기 처리')
    _box(ax, 6.5, 0.6, 2.0, 1.0, 'MySQL RDS', 'Source of Truth', filled=True)

    _arrow(ax, 4.0, 3.6, 1.5, 1.6, 'Fire & Forget', dashed=True)
    _arrow(ax, 2.5, 1.1, 3.5, 1.1, '폴링')
    _arrow(ax, 5.5, 1.1, 6.5, 1.1, 'DB 저장')

    # Scheduler
    _box(ax, 9.2, 0.8, 2.3, 0.7, 'Scheduler (5분)', '정합성 검증')
    _arrow(ax, 9.2, 1.3, 8.5, 1.1, '', dashed=True)
    _arrow(ax, 10.0, 1.5, 7.8, 3.6, '', dashed=True)

    fig.tight_layout()
    return fig_to_bytes(fig)


def create_reversal():
    fig, ax = plt.subplots(figsize=(8.5, 3.8))
    fig.set_facecolor(WHITE)

    phases = ['Phase 6\n비관적 락', 'Phase 10\nSQS 비동기',
              'Phase 10-2\nRedis 중복체크', 'Phase 10-3\nPooling']
    values = [1730, 516, 1300, 658]
    # Accent only the reversal bar — rest is gray
    cols = ['#94A3B8', '#64748B', INK, ACCENT]

    bars = ax.bar(phases, values, color=cols, width=0.5, edgecolor=WHITE, lw=2)
    for bar,v,c in zip(bars, values, cols):
        ax.text(bar.get_x()+bar.get_width()/2, bar.get_height()+30,
                f'{v:,}ms', ha='center', fontsize=11, fontweight='bold',
                color=INK, fontfamily=FONT)

    # Highlight reversal
    ax.annotate('성능 역전', xy=(2, 1300), xytext=(2.7, 1550),
                fontsize=12, fontweight='bold', color=INK, fontfamily=FONT,
                arrowprops=dict(arrowstyle='->', color=INK, lw=1.8))

    # Dashed box around reversal
    ax.add_patch(plt.Rectangle((1.55, 0), 0.9, 1350, fill=False,
                               edgecolor=INK, lw=1.5, ls='--'))

    ax.set_title('투표 API p95 변화 추이', fontsize=12, fontweight='bold',
                 color=INK, pad=12, fontfamily=FONT)
    ax.set_ylabel('p95 (ms)', fontsize=9, color=DIM, fontfamily=FONT)
    for s in ['top','right']: ax.spines[s].set_visible(False)
    for s in ['left','bottom']: ax.spines[s].set_color(LINE)
    ax.tick_params(colors=DIM, labelsize=8)
    for l in ax.get_xticklabels()+ax.get_yticklabels(): l.set_fontfamily(FONT)
    fig.tight_layout()
    return fig_to_bytes(fig)


def create_connection_model():
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4.2))
    fig.set_facecolor(WHITE)

    # ── Left: Lettuce ──
    ax1.set_xlim(0,10); ax1.set_ylim(0,10); ax1.axis('off')
    ax1.set_title('Lettuce — 단일 커넥션 멀티플렉싱', fontsize=10, fontweight='bold',
                  color=INK, fontfamily=FONT, pad=8)

    for i in range(5):
        y = 8.5 - i*1.2
        _box(ax1, 0.2, y-0.2, 1.6, 0.5, f'요청 {i+1}')
        ax1.annotate('', xy=(4.3, 5.05), xytext=(1.8, y),
                     arrowprops=dict(arrowstyle='->', color=LINE, lw=0.8))

    # Bottleneck funnel
    ax1.add_patch(FancyBboxPatch((4.2, 4.2), 2.2, 1.6, boxstyle="round,pad=0.08",
                                facecolor=INK, edgecolor=INK, lw=1.5))
    ax1.text(5.3, 5.25, 'TCP 1개', fontsize=11, fontweight='bold',
             color=WHITE, ha='center', fontfamily=FONT)
    ax1.text(5.3, 4.7, '직렬화 병목', fontsize=8,
             color=MUTE, ha='center', fontfamily=FONT)

    _box(ax1, 7.2, 4.4, 2.2, 1.2, 'Redis', '', filled=True)
    _arrow(ax1, 6.4, 5.0, 7.2, 5.0)

    ax1.text(5, 1.5, 'p95: 1,300ms', fontsize=13, fontweight='bold',
             color=INK, ha='center', fontfamily=FONT)
    ax1.add_patch(FancyBboxPatch((2.8, 1.0), 4.4, 0.8, boxstyle="round,pad=0.08",
                                facecolor=BG, edgecolor=LINE, lw=1))

    # ── Right: HikariCP ──
    ax2.set_xlim(0,10); ax2.set_ylim(0,10); ax2.axis('off')
    ax2.set_title('HikariCP — 커넥션 풀 병렬 처리', fontsize=10, fontweight='bold',
                  color=INK, fontfamily=FONT, pad=8)

    for i in range(5):
        y = 8.5 - i*1.2
        _box(ax2, 0.2, y-0.2, 1.6, 0.5, f'요청 {i+1}')

    # Multiple connections
    conn_ys = [7.8, 6.2, 4.6, 3.0]
    ax2.add_patch(FancyBboxPatch((3.3, 2.2), 2.6, 6.5, boxstyle="round,pad=0.12",
                                facecolor='none', edgecolor=ACCENT, lw=1.5, ls='--'))
    ax2.text(4.6, 2.5, '30개 커넥션 풀', fontsize=8, fontweight='bold',
             color=ACCENT, ha='center', fontfamily=FONT)

    for i, cy in enumerate(conn_ys):
        ax2.add_patch(FancyBboxPatch((3.6, cy-0.2), 1.8, 0.5, boxstyle="round,pad=0.04",
                                    facecolor=ACCENT, edgecolor=ACCENT, lw=1))
        ax2.text(4.5, cy+0.05, f'Conn {i+1}', fontsize=7.5, fontweight='bold',
                 color=WHITE, ha='center', fontfamily=FONT)
        _arrow(ax2, 5.4, cy+0.05, 7.0, 5.05)

    for i in range(5):
        y = 8.5 - i*1.2
        t = conn_ys[min(i,3)]
        ax2.annotate('', xy=(3.6, t+0.05), xytext=(1.8, y),
                     arrowprops=dict(arrowstyle='->', color=ACC_L, lw=0.8))

    _box(ax2, 7.0, 4.2, 2.2, 1.6, 'MySQL', '병렬 처리', filled=True)

    ax2.text(5, 1.5, 'p95: 516ms', fontsize=13, fontweight='bold',
             color=ACCENT, ha='center', fontfamily=FONT)
    ax2.add_patch(FancyBboxPatch((2.8, 1.0), 4.4, 0.8, boxstyle="round,pad=0.08",
                                facecolor=ACC_L, edgecolor=ACCENT, lw=1))

    fig.tight_layout(w_pad=2)
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════
#  DOCX HELPERS
# ═══════════════════════════════════════════

def hex_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[:2],16), int(h[2:4],16), int(h[4:6],16))

def _shd(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color.lstrip("#")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def _run(p, text, sz=9.5, bold=False, color=INK, italic=False):
    r = p.add_run(text)
    r.font.size = Pt(sz)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = hex_rgb(color)
    r.font.name = FONT
    rPr = r._element.get_or_add_rPr()
    rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT}"/>'))
    return r

def para(doc, text='', sz=9.5, bold=False, color=INK, align=WD_ALIGN_PARAGRAPH.LEFT,
         sb=0, sa=4, italic=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    if text:
        _run(p, text, sz, bold, color, italic)
    return p

def section_head(doc, num, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(2)

    r_num = _run(p, f' {num} ', 13, True, WHITE)
    rPr = r_num._element.get_or_add_rPr()
    rPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{ACCENT.lstrip("#")}"/>'))
    _run(p, '  ', 13)
    _run(p, title, 15, True, INK)

    para(doc, subtitle, 9.5, False, DIM, sb=0, sa=10)

def label(doc, text, color=ACCENT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    _run(p, f'— {text}', 9, True, color)

def insight(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.4)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="16" w:space="8" w:color="{ACCENT.lstrip("#")}"/>'
        f'</w:pBdr>'))
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{ACC_L.lstrip("#")}"/>'))
    _run(p, 'INSIGHT  ', 7.5, True, ACCENT)
    _run(p, text, 9, False, INK)

def table(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Clean borders
    tbl = t._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblPr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="{LINE.lstrip("#")}"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="{LINE.lstrip("#")}"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="{LINE.lstrip("#")}"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'))
    # Header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        _shd(c, NAVY)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, h, 8.5, True, WHITE)
    # Rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            if ri % 2 == 1: _shd(c, BG)
            p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, str(val), 8.5, False, INK)
    return t

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._element.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="{LINE.lstrip("#")}"/>'
        f'</w:pBdr>'))


# ═══════════════════════════════════════════
#  BUILD
# ═══════════════════════════════════════════

def build():
    doc = Document()

    # Page setup
    sec = doc.sections[0]
    sec.page_width = Cm(21); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)

    # Default font
    style = doc.styles['Normal']
    style.font.name = FONT; style.font.size = Pt(9.5)
    style.font.color.rgb = hex_rgb(INK)
    rPr = style.element.get_or_add_rPr()
    rPr.append(parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{FONT}"/>'))

    # ── HERO ──
    doc.add_picture(create_hero(), width=Cm(17))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tech pills
    p = para(doc, '', sa=1, sb=4, align=WD_ALIGN_PARAGRAPH.CENTER)
    techs = ['Java 21','Spring Boot 3.5','MySQL 8.0','Redis 7','AWS SQS',
             'QueryDSL','Docker','Terraform','k6','Prometheus']
    for i, t in enumerate(techs):
        _run(p, t, 8, False, ACCENT)
        if i < len(techs)-1:
            _run(p, '  ·  ', 8, False, LINE)

    para(doc, '반려동물 사진 콘테스트 플랫폼 — 챌린지별 투표 → 실시간 랭킹',
         9, False, DIM, WD_ALIGN_PARAGRAPH.CENTER, 2, 2)
    para(doc, '테스트 데이터 10K 회원 · 100K 엔트리 · 1M 투표  |  EC2 t3.small · RDS db.t3.micro',
         8, False, MUTE, WD_ALIGN_PARAGRAPH.CENTER, 0, 6)

    divider(doc)

    # ═══════════════════════════════════════
    #  01. DB QUERY OPTIMIZATION
    # ═══════════════════════════════════════
    section_head(doc, '01', 'DB 쿼리 최적화',
                 'Ranking API p95 응답시간 4.76s → 146ms (97% 감소)')

    label(doc, 'PROBLEM')
    para(doc, 'k6 부하 테스트(50 VUs)에서 랭킹 API p95가 4.76초로 측정되었다. '
         'Grafana + p6spy 분석 결과 세 가지 병목을 확인했다.', sb=2)

    table(doc,
        ['병목', '원인', '영향'],
        [['데이터 전송량', 'Java Stream.limit()으로 애플리케이션에서 자름', '불필요한 네트워크/메모리 사용'],
         ['N+1 쿼리', '연관 엔티티 개별 SELECT — 요청당 21개 쿼리', 'DB 커넥션 점유 시간 증가'],
         ['filesort', 'vote_count 정렬에 인덱스 없음', 'CPU 사용률 증가']])

    label(doc, 'SOLUTION')
    para(doc, '상위 레이어부터 순차 최적화 — 상위 병목을 먼저 제거해야 하위 레이어의 정확한 개선 효과를 측정할 수 있다.',
         9.5, True, sb=2)

    doc.add_picture(create_waterfall(), width=Cm(14.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table(doc,
        ['단계', '최적화', 'Before', 'After', '개선율'],
        [['1단계', 'Pageable — DB LIMIT 적용', 'p95 4,760ms', 'p95 884ms', '81%↓'],
         ['2단계', 'Fetch Join — N+1 해결 (21→1 쿼리)', 'p95 884ms', 'p95 234ms', '73%↓'],
         ['3단계', '복합 인덱스 (challenge_id, vote_count DESC)', 'p95 234ms', 'p95 146ms', '38%↓']])

    doc.add_picture(create_query_throughput(), width=Cm(14.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    label(doc, 'RESULT')
    table(doc,
        ['지표', '최적화 전 (50 VUs)', '최적화 후 (300 VUs)', '개선'],
        [['랭킹 API p95', '4,760ms', '178ms', '97% 감소'],
         ['요청당 쿼리 수', '21개', '1개', '95% 감소'],
         ['전체 처리량', '10.9 RPS', '308 RPS', '28배 증가'],
         ['에러율', '-', '0%', '300 VUs 안정']])

    insight(doc, '최적화 순서가 중요하다. 1단계에서 데이터 전송량을 줄이지 않으면 '
            '2단계 Fetch Join의 효과가 과대/과소 측정될 수 있다. '
            '측정 가능한 최적화를 위해 상위 레이어부터 제거하는 원칙을 적용했다.')

    divider(doc)

    # ═══════════════════════════════════════
    #  02. CONCURRENCY CONTROL
    # ═══════════════════════════════════════
    section_head(doc, '02', '동시성 제어',
                 '3가지 전략 비교 분석 후 비관적 락 채택')

    label(doc, 'PROBLEM')
    para(doc, '투표 기능에서 두 가지 레이스 컨디션이 발생했다.', sb=2)

    table(doc,
        ['레이스 컨디션', '원인', '결과'],
        [['Check-then-Act', 'SELECT(중복 확인)과 INSERT(투표 생성) 사이 갭', '중복 투표 허용'],
         ['Lost Update', '동시 voteCount++ — 한쪽 증가분 유실', '투표 수 불일치']])

    para(doc, 'synchronized는 JVM 레벨 락 → 다중 인스턴스 무효. '
         '@Transactional 단독으로는 MySQL REPEATABLE READ에서 Lost Update 미방지.',
         9, False, DIM, sb=4)

    label(doc, 'SOLUTION')
    para(doc, '동일 조건(50명 동시 투표, 동일 Entry)에서 3가지 전략을 구현하고 테스트했다.', sb=2)

    doc.add_picture(create_concurrency(), width=Cm(14.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table(doc,
        ['전략', '성공률', '소요시간', '문제점'],
        [['비관적 락 (채택)', '100%', '1,730ms', '없음'],
         ['낙관적 락 (@Version)', '16%', '885ms', '42/50 실패 — version NULL 이슈'],
         ['Atomic + Retry', '94%', '2,980ms', 'Retry Storm — 재시도가 재충돌 유발']])

    doc.add_picture(create_retry_vs_pessimistic(), width=Cm(15.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    label(doc, 'TROUBLESHOOTING')
    table(doc,
        ['문제', '원인', '해결'],
        [['@Version NULL', '기존 DB 데이터에 version 값 없음', 'ALTER TABLE 기본값 마이그레이션'],
         ['@Transactional 미적용', '동일 클래스 내부 호출 시 AOP 프록시 우회', '@Lazy 자기 주입으로 프록시 호출 보장']])

    insight(doc, '"느린 비관적 락이 역설적으로 빠르다." '
            '고경쟁(hot-spot) 환경에서 낙관적 접근은 충돌 → 재시도 → 재충돌의 '
            'Retry Storm을 일으켜 전체 응답시간이 증가한다. 경쟁 강도에 따라 전략을 선택해야 한다.')

    divider(doc)

    # ═══════════════════════════════════════
    #  03. ASYNC ARCHITECTURE
    # ═══════════════════════════════════════
    section_head(doc, '03', '비동기 투표 아키텍처',
                 '"빠른 기술 ≠ 빠른 시스템" — 커넥션 모델의 중요성')

    label(doc, 'PROBLEM')
    para(doc, '비관적 락으로 정합성은 확보했지만 p95 1.73초 — 동기 DB 쓰기가 병목.', sb=2)

    label(doc, 'SOLUTION — 읽기/쓰기 분리 (CQRS)')

    doc.add_picture(create_architecture(), width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    para(doc, 'Fast Path: 투표 → Redis ZINCRBY → 사용자 즉시 응답 (~5ms)\n'
         'Slow Path: SQS 메시지 발행 → Consumer 비동기 DB 영속화\n'
         '정합성: 5분 Scheduler가 Redis ↔ DB 검증 및 보정', sb=4)

    label(doc, 'MQ 선택 — SQS')
    table(doc,
        ['기준', 'Kafka (MSK)', 'RabbitMQ', 'SQS (채택)'],
        [['월 비용', '$200+', '$50-100', '$1'],
         ['운영', 'ZooKeeper', 'Erlang', 'Zero-Ops'],
         ['확장', '수동 파티션', '수동', '자동']])

    para(doc, '순서 보장 불필요 + 멱등성 설계(Application 중복 체크 + DB Unique Constraint) '
         '→ SQS Standard Queue. DLQ(3회 재시도 → 14일 보관)로 메시지 유실 방지.',
         9, False, DIM, sb=4)

    label(doc, 'DISCOVERY — 성능 역전 현상')
    para(doc, '중복 투표 체크를 DB → Redis(SISMEMBER)로 이관 → 오히려 2.5배 느려짐.',
         9.5, True, sb=4)

    doc.add_picture(create_reversal(), width=Cm(14.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    label(doc, 'ROOT CAUSE — 커넥션 모델의 차이')

    doc.add_picture(create_connection_model(), width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    table(doc,
        ['', 'Lettuce (Redis)', 'HikariCP (MySQL)'],
        [['커넥션', '단일 TCP 멀티플렉싱', '30개 커넥션 풀'],
         ['50 동시 요청', '1개 파이프 직렬화 → 병목', '30개 병렬 분산'],
         ['결과', 'p95 1,300ms', 'p95 516ms']])

    label(doc, 'RESOLUTION')
    para(doc, 'Lettuce Connection Pooling(commons-pool2, max-active:20) + '
         'Redis Pipeline(ZINCRBY+ZSCORE 단일 RTT) 적용. '
         '최종적으로 DB 중복 체크를 유지하는 Hybrid 전략 채택.', sb=4)

    label(doc, 'RESULT')
    table(doc,
        ['단계', 'p95', '설명'],
        [['Phase 6 — 비관적 락', '1,730ms', '동기 DB 쓰기 병목'],
         ['Phase 10 — SQS 비동기', '516ms', 'DB 쓰기 분리 (70%↓)'],
         ['Phase 10-2 — Redis 중복체크', '1,300ms', '성능 역전 (Lettuce 단일 커넥션)'],
         ['Phase 10-3 — Pooling+Pipeline', '658ms', '최종 (비관적 락 대비 62%↓)']])

    insight(doc, '"빠른 기술 ≠ 빠른 시스템." '
            'Redis O(1) 연산보다 HikariCP 30개 커넥션의 병렬 처리가 50 동시 사용자 환경에서 더 효과적이었다. '
            '개별 연산 속도가 아닌 커넥션 모델이 실제 성능을 결정한다.')

    divider(doc)

    # ── Tech decisions summary ──
    para(doc, '기술적 의사결정 요약', 13, True, INK, sb=16, sa=8)

    table(doc,
        ['영역', '선택', '근거'],
        [['캐시', 'Redis (Global Cache)', 'Scale-out 시 로컬 캐시 일관성 보장 불가'],
         ['동시성', '비관적 락 (SELECT FOR UPDATE)', 'Hot-spot 100% 정합성. 낙관적 락은 Retry Storm'],
         ['MQ', 'SQS Standard Queue', '순서 불필요 + 멱등성 설계. $1/mo, Zero-Ops'],
         ['중복 체크', 'DB (Hybrid)', 'Lettuce 단일 커넥션 병목. DB 풀이 동시성에서 유리'],
         ['데이터 정합성', 'DB = Source of Truth', 'Redis 장애 시 DB 복구 가능. 역방향 불가']])

    # Save
    out = os.path.join(OUT_DIR, 'PetStar_Portfolio_v2.docx')
    doc.save(out)
    print(f'Saved: {out}')


if __name__ == '__main__':
    build()
