이력서에 #!/usr/bin/env python3
"""PetStar Portfolio DOCX Generator v1"""

import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.patches as patches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import numpy as np
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── Colors ──
C_PRIMARY   = '#1A2744'   # Dark navy
C_ACCENT    = '#2563EB'   # Blue
C_ACCENT2   = '#0EA5E9'   # Cyan
C_SUCCESS   = '#059669'   # Green
C_DANGER    = '#DC2626'   # Red
C_WARNING   = '#D97706'   # Amber
C_GRAY      = '#6B7280'   # Gray
C_LIGHT_BG  = '#F1F5F9'   # Light slate
C_WHITE     = '#FFFFFF'
C_DARK      = '#111827'

# ── Font Setup ──
FONT_KR = 'Apple SD Gothic Neo'
plt.rcParams['font.family'] = FONT_KR
plt.rcParams['axes.unicode_minus'] = False

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
IMG_DIR = os.path.join(OUT_DIR, '_img_temp')
os.makedirs(IMG_DIR, exist_ok=True)


# ═══════════════════════════════════════════
#  DIAGRAM GENERATORS
# ═══════════════════════════════════════════

def fig_to_bytes(fig, dpi=200):
    buf = BytesIO()
    fig.savefig(buf, format='png', dpi=dpi, bbox_inches='tight',
                facecolor=fig.get_facecolor(), edgecolor='none')
    plt.close(fig)
    buf.seek(0)
    return buf


def create_hero_banner():
    """Project summary hero banner"""
    fig, ax = plt.subplots(figsize=(10, 3.2))
    fig.set_facecolor(C_PRIMARY)
    ax.set_facecolor(C_PRIMARY)
    ax.set_xlim(0, 10); ax.set_ylim(0, 3.2)
    ax.axis('off')

    # Title
    ax.text(5, 2.55, 'PetStar', fontsize=36, fontweight='bold',
            color=C_WHITE, ha='center', va='center', fontfamily=FONT_KR)
    ax.text(5, 1.95, '반려동물 사진 콘테스트 플랫폼  |  백엔드 성능 최적화 프로젝트',
            fontsize=12, color='#94A3B8', ha='center', va='center', fontfamily=FONT_KR)

    # Metric boxes
    metrics = [
        ('응답시간', '97%↓', '4.76s → 146ms'),
        ('처리량', '28x↑', '10.9 → 308 RPS'),
        ('에러율', '0%', '300 VUs 안정'),
    ]
    box_w, box_h = 2.4, 0.85
    start_x = 5 - (len(metrics) * box_w + (len(metrics)-1)*0.3) / 2

    for i, (label, value, detail) in enumerate(metrics):
        x = start_x + i * (box_w + 0.3)
        rect = FancyBboxPatch((x, 0.25), box_w, box_h, boxstyle="round,pad=0.08",
                              facecolor='#1E3A5F', edgecolor='#334155', linewidth=1)
        ax.add_patch(rect)
        ax.text(x + box_w/2, 0.82, value, fontsize=18, fontweight='bold',
                color=C_ACCENT2, ha='center', va='center', fontfamily=FONT_KR)
        ax.text(x + box_w/2, 0.53, f'{label}: {detail}', fontsize=7.5,
                color='#CBD5E1', ha='center', va='center', fontfamily=FONT_KR)

    return fig_to_bytes(fig)


def create_optimization_waterfall():
    """3-step optimization waterfall chart"""
    fig, ax = plt.subplots(figsize=(9, 4))
    fig.set_facecolor(C_WHITE)
    ax.set_facecolor(C_WHITE)

    stages = ['Before', '1단계\nPageable', '2단계\nFetch Join', '3단계\n복합 인덱스']
    values = [4760, 884, 234, 146]
    colors = ['#DC2626', '#F59E0B', '#3B82F6', '#059669']
    reductions = ['', '81%↓', '73%↓', '38%↓']

    bars = ax.bar(stages, values, color=colors, width=0.55, edgecolor='white', linewidth=2)

    for i, (bar, val, red) in enumerate(zip(bars, values, reductions)):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 80,
                f'{val:,}ms', ha='center', va='bottom', fontsize=12,
                fontweight='bold', color=colors[i], fontfamily=FONT_KR)
        if red:
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height()/2,
                    red, ha='center', va='center', fontsize=14,
                    fontweight='bold', color='white', fontfamily=FONT_KR)

    # Arrow from first to last
    ax.annotate('', xy=(3, 300), xytext=(0, 4500),
                arrowprops=dict(arrowstyle='->', color=C_ACCENT, lw=2.5,
                                connectionstyle='arc3,rad=-0.3'))
    ax.text(2.1, 3200, '97% 감소', fontsize=14, fontweight='bold',
            color=C_ACCENT, ha='center', fontfamily=FONT_KR, rotation=-30)

    ax.set_ylabel('p95 응답시간 (ms)', fontsize=11, color=C_DARK, fontfamily=FONT_KR)
    ax.set_title('Ranking API p95 응답시간 최적화 과정', fontsize=14,
                 fontweight='bold', color=C_DARK, pad=15, fontfamily=FONT_KR)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#E5E7EB')
    ax.spines['bottom'].set_color('#E5E7EB')
    ax.tick_params(colors=C_GRAY)
    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontfamily(FONT_KR)

    fig.tight_layout()
    return fig_to_bytes(fig)


def create_query_comparison():
    """N+1 query count comparison"""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9, 3.2))
    fig.set_facecolor(C_WHITE)

    # Left: Query count
    cats = ['Before\n(N+1)', 'After\n(Fetch Join)']
    vals = [21, 1]
    colors = [C_DANGER, C_SUCCESS]
    bars = ax1.bar(cats, vals, color=colors, width=0.5, edgecolor='white', linewidth=2)
    for bar, val in zip(bars, vals):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5,
                 f'{val}개', ha='center', fontsize=14, fontweight='bold',
                 color=C_DARK, fontfamily=FONT_KR)
    ax1.set_title('요청당 쿼리 수', fontsize=13, fontweight='bold',
                  color=C_DARK, fontfamily=FONT_KR, pad=10)
    ax1.set_ylabel('쿼리 수', fontsize=10, color=C_GRAY, fontfamily=FONT_KR)
    ax1.spines['top'].set_visible(False)
    ax1.spines['right'].set_visible(False)
    ax1.spines['left'].set_color('#E5E7EB')
    ax1.spines['bottom'].set_color('#E5E7EB')
    for label in ax1.get_xticklabels() + ax1.get_yticklabels():
        label.set_fontfamily(FONT_KR)

    # Right: Throughput
    cats2 = ['Before\n(50 VUs)', 'After\n(300 VUs)']
    vals2 = [10.9, 308]
    bars2 = ax2.bar(cats2, vals2, color=[C_DANGER, C_SUCCESS], width=0.5,
                    edgecolor='white', linewidth=2)
    for bar, val in zip(bars2, vals2):
        ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 5,
                 f'{val} RPS', ha='center', fontsize=14, fontweight='bold',
                 color=C_DARK, fontfamily=FONT_KR)
    ax2.set_title('처리량 (Requests/sec)', fontsize=13, fontweight='bold',
                  color=C_DARK, fontfamily=FONT_KR, pad=10)
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.spines['left'].set_color('#E5E7EB')
    ax2.spines['bottom'].set_color('#E5E7EB')
    for label in ax2.get_xticklabels() + ax2.get_yticklabels():
        label.set_fontfamily(FONT_KR)

    fig.tight_layout(w_pad=3)
    return fig_to_bytes(fig)


def create_concurrency_comparison():
    """Concurrency strategy comparison chart"""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9, 3.5))
    fig.set_facecolor(C_WHITE)

    strategies = ['비관적 락\n(채택)', '낙관적 락', 'Atomic\n+ Retry']

    # Success rate
    rates = [100, 16, 94]
    colors_rate = [C_SUCCESS, C_DANGER, C_WARNING]
    bars1 = ax1.bar(strategies, rates, color=colors_rate, width=0.5,
                    edgecolor='white', linewidth=2)
    for bar, val in zip(bars1, rates):
        ax1.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1.5,
                 f'{val}%', ha='center', fontsize=14, fontweight='bold',
                 color=C_DARK, fontfamily=FONT_KR)
    ax1.set_ylim(0, 115)
    ax1.axhline(y=100, color='#E5E7EB', linestyle='--', linewidth=1)
    ax1.set_title('정합성 (50명 동시 투표)', fontsize=13, fontweight='bold',
                  color=C_DARK, fontfamily=FONT_KR, pad=10)
    ax1.set_ylabel('성공률 (%)', fontsize=10, color=C_GRAY, fontfamily=FONT_KR)
    ax1.spines['top'].set_visible(False)
    ax1.spines['right'].set_visible(False)
    ax1.spines['left'].set_color('#E5E7EB')
    ax1.spines['bottom'].set_color('#E5E7EB')
    for label in ax1.get_xticklabels() + ax1.get_yticklabels():
        label.set_fontfamily(FONT_KR)

    # Response time
    times = [1730, 885, 2980]
    colors_time = [C_ACCENT, C_DANGER, C_WARNING]
    bars2 = ax2.bar(strategies, times, color=colors_time, width=0.5,
                    edgecolor='white', linewidth=2)
    for bar, val in zip(bars2, times):
        ax2.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 50,
                 f'{val:,}ms', ha='center', fontsize=12, fontweight='bold',
                 color=C_DARK, fontfamily=FONT_KR)
    ax2.set_title('소요시간 비교', fontsize=13, fontweight='bold',
                  color=C_DARK, fontfamily=FONT_KR, pad=10)
    ax2.set_ylabel('소요시간 (ms)', fontsize=10, color=C_GRAY, fontfamily=FONT_KR)
    ax2.spines['top'].set_visible(False)
    ax2.spines['right'].set_visible(False)
    ax2.spines['left'].set_color('#E5E7EB')
    ax2.spines['bottom'].set_color('#E5E7EB')
    for label in ax2.get_xticklabels() + ax2.get_yticklabels():
        label.set_fontfamily(FONT_KR)

    fig.tight_layout(w_pad=3)
    return fig_to_bytes(fig)


def create_retry_storm_diagram():
    """Visual explanation of retry storm vs pessimistic lock"""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9, 4))
    fig.set_facecolor(C_WHITE)

    # ── Left: Optimistic Lock - Retry Storm ──
    ax1.set_xlim(0, 10); ax1.set_ylim(0, 10)
    ax1.axis('off')
    ax1.set_title('낙관적 락: Retry Storm', fontsize=13, fontweight='bold',
                  color=C_DANGER, fontfamily=FONT_KR, pad=10)

    # Timeline
    users_y = [8.5, 7.0, 5.5, 4.0]
    user_labels = ['User 1', 'User 2', 'User 3', 'User N']
    for i, (y, label) in enumerate(zip(users_y, user_labels)):
        ax1.add_patch(FancyBboxPatch((0.3, y-0.25), 1.4, 0.5, boxstyle="round,pad=0.05",
                                    facecolor='#DBEAFE', edgecolor=C_ACCENT, linewidth=1))
        ax1.text(1.0, y, label, ha='center', va='center', fontsize=8,
                 color=C_ACCENT, fontweight='bold', fontfamily=FONT_KR)

    # Arrows showing retries
    for y in users_y[:3]:
        # First attempt - fail
        ax1.annotate('', xy=(4.5, y), xytext=(1.8, y),
                     arrowprops=dict(arrowstyle='->', color=C_DANGER, lw=1.2))
        ax1.text(3.2, y+0.2, '충돌!', fontsize=7, color=C_DANGER,
                 ha='center', fontfamily=FONT_KR)
        # Retry
        ax1.annotate('', xy=(7.0, y-0.3), xytext=(4.8, y-0.3),
                     arrowprops=dict(arrowstyle='->', color=C_WARNING, lw=1.2, linestyle='--'))
        ax1.text(5.9, y-0.55, '재시도', fontsize=7, color=C_WARNING,
                 ha='center', fontfamily=FONT_KR)
        # Re-retry
        ax1.annotate('', xy=(9.5, y-0.6), xytext=(7.3, y-0.6),
                     arrowprops=dict(arrowstyle='->', color='#991B1B', lw=1.2, linestyle='--'))

    # Explosion
    ax1.text(5, 1.8, '충돌 → 재시도 → 재충돌', fontsize=10, fontweight='bold',
             color=C_DANGER, ha='center', fontfamily=FONT_KR)
    ax1.text(5, 1.0, '지수적 경쟁 증가', fontsize=9,
             color='#991B1B', ha='center', fontfamily=FONT_KR)

    # ── Right: Pessimistic Lock - Ordered Queue ──
    ax2.set_xlim(0, 10); ax2.set_ylim(0, 10)
    ax2.axis('off')
    ax2.set_title('비관적 락: 순차 처리', fontsize=13, fontweight='bold',
                  color=C_SUCCESS, fontfamily=FONT_KR, pad=10)

    # Queue visualization
    queue_y = 7.5
    ax2.add_patch(FancyBboxPatch((0.5, queue_y-0.6), 4, 1.2, boxstyle="round,pad=0.1",
                                facecolor='#DCFCE7', edgecolor=C_SUCCESS, linewidth=1.5))
    ax2.text(2.5, queue_y, '대기열 (SELECT FOR UPDATE)', fontsize=9,
             ha='center', va='center', fontweight='bold', color=C_SUCCESS, fontfamily=FONT_KR)

    # Sequential processing
    for i, (label, y) in enumerate(zip(['User 1', 'User 2', 'User 3', 'User N'],
                                        [5.8, 4.5, 3.2, 1.9])):
        color_fill = C_SUCCESS if i == 0 else '#E5E7EB'
        color_text = C_WHITE if i == 0 else C_GRAY
        ax2.add_patch(FancyBboxPatch((1, y-0.25), 3, 0.5, boxstyle="round,pad=0.05",
                                    facecolor=color_fill, edgecolor=C_SUCCESS, linewidth=1))
        ax2.text(2.5, y, f'{label} → 처리 완료' if i == 0 else f'{label} → 대기 중',
                 ha='center', va='center', fontsize=8, color=color_text,
                 fontweight='bold', fontfamily=FONT_KR)
        if i < 3:
            ax2.annotate('', xy=(2.5, y-0.35), xytext=(2.5, y-0.65),
                         arrowprops=dict(arrowstyle='->', color=C_SUCCESS, lw=1.2))

    # Result labels
    ax2.add_patch(FancyBboxPatch((5.5, 4.0), 4, 2.5, boxstyle="round,pad=0.15",
                                facecolor='#F0FDF4', edgecolor=C_SUCCESS, linewidth=1.5))
    ax2.text(7.5, 5.7, '충돌 없음', fontsize=11, fontweight='bold',
             color=C_SUCCESS, ha='center', fontfamily=FONT_KR)
    ax2.text(7.5, 5.0, '100% 성공', fontsize=11, fontweight='bold',
             color=C_SUCCESS, ha='center', fontfamily=FONT_KR)
    ax2.text(7.5, 4.4, '예측 가능한 성능', fontsize=9,
             color=C_GRAY, ha='center', fontfamily=FONT_KR)

    fig.tight_layout(w_pad=2)
    return fig_to_bytes(fig)


def create_async_architecture():
    """Async vote system architecture diagram"""
    fig, ax = plt.subplots(figsize=(10, 5.5))
    fig.set_facecolor(C_WHITE)
    ax.set_xlim(0, 12); ax.set_ylim(0, 7)
    ax.axis('off')

    ax.text(6, 6.6, '비동기 투표 시스템 아키텍처', fontsize=16, fontweight='bold',
            color=C_PRIMARY, ha='center', fontfamily=FONT_KR)

    def draw_box(ax, x, y, w, h, label, sublabel='', color=C_ACCENT, text_color=C_WHITE):
        ax.add_patch(FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.08",
                                   facecolor=color, edgecolor=color, linewidth=1.5))
        ax.text(x+w/2, y+h/2 + (0.12 if sublabel else 0), label,
                ha='center', va='center', fontsize=10, fontweight='bold',
                color=text_color, fontfamily=FONT_KR)
        if sublabel:
            ax.text(x+w/2, y+h/2-0.18, sublabel, ha='center', va='center',
                    fontsize=7.5, color=text_color, alpha=0.8, fontfamily=FONT_KR)

    def draw_arrow(ax, x1, y1, x2, y2, label='', color=C_GRAY, style='->', lw=1.8):
        ax.annotate('', xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle=style, color=color, lw=lw))
        if label:
            mx, my = (x1+x2)/2, (y1+y2)/2
            ax.text(mx, my+0.18, label, ha='center', va='center', fontsize=7.5,
                    color=color, fontweight='bold', fontfamily=FONT_KR,
                    bbox=dict(boxstyle='round,pad=0.15', facecolor=C_WHITE, edgecolor='none'))

    # ── Fast Path (top) ──
    ax.add_patch(FancyBboxPatch((0.2, 3.5), 11.6, 2.7, boxstyle="round,pad=0.15",
                                facecolor='#EFF6FF', edgecolor='#BFDBFE', linewidth=1.5,
                                linestyle='--'))
    ax.text(6, 6.0, '빠른 경로 (Fast Path) — 즉시 응답', fontsize=10,
            color=C_ACCENT, ha='center', fontweight='bold', fontfamily=FONT_KR)

    draw_box(ax, 0.5, 4.0, 1.6, 1.2, '클라이언트', '투표 요청', '#374151')
    draw_box(ax, 3.0, 4.0, 2.0, 1.2, 'Spring Boot', 'API Server', C_ACCENT)
    draw_box(ax, 6.2, 4.0, 2.0, 1.2, 'Redis', 'Sorted Set', '#DC2626')

    draw_arrow(ax, 2.1, 4.6, 3.0, 4.6, '투표', C_DARK)
    draw_arrow(ax, 5.0, 4.6, 6.2, 4.6, 'ZINCRBY', '#DC2626')
    draw_arrow(ax, 6.2, 4.3, 5.0, 4.3, '즉시 반영', '#DC2626')
    draw_arrow(ax, 3.0, 4.3, 2.1, 4.3, '응답', C_SUCCESS)

    # Response time badge
    ax.add_patch(FancyBboxPatch((9.0, 4.2), 2.5, 0.8, boxstyle="round,pad=0.1",
                                facecolor=C_SUCCESS, edgecolor=C_SUCCESS))
    ax.text(10.25, 4.6, '~5ms 응답', fontsize=12, fontweight='bold',
            color=C_WHITE, ha='center', va='center', fontfamily=FONT_KR)

    # ── Slow Path (bottom) ──
    ax.add_patch(FancyBboxPatch((0.2, 0.3), 11.6, 2.7, boxstyle="round,pad=0.15",
                                facecolor='#FFF7ED', edgecolor='#FED7AA', linewidth=1.5,
                                linestyle='--'))
    ax.text(6, 2.8, '느린 경로 (Slow Path) — 비동기 영속화', fontsize=10,
            color=C_WARNING, ha='center', fontweight='bold', fontfamily=FONT_KR)

    draw_box(ax, 0.5, 0.7, 2.0, 1.2, 'SQS Queue', '메시지 발행', '#D97706')
    draw_box(ax, 3.5, 0.7, 2.0, 1.2, 'Consumer', '비동기 처리', '#7C3AED')
    draw_box(ax, 6.5, 0.7, 2.0, 1.2, 'MySQL RDS', 'Source of Truth', C_ACCENT)

    draw_arrow(ax, 4.0, 4.0, 1.5, 1.9, 'Fire & Forget', C_WARNING)
    draw_arrow(ax, 2.5, 1.3, 3.5, 1.3, '폴링', '#7C3AED')
    draw_arrow(ax, 5.5, 1.3, 6.5, 1.3, 'DB 저장', C_ACCENT)

    # DLQ
    draw_box(ax, 3.5, -0.1, 2.0, 0.5, 'DLQ (3회 재시도 → 14일 보관)', '', '#991B1B', C_WHITE)

    # Scheduler
    draw_box(ax, 9.2, 1.0, 2.3, 0.8, 'Scheduler', '5분 정합성 검증', '#6B7280')
    draw_arrow(ax, 9.2, 1.6, 8.5, 1.3, '', C_GRAY)
    draw_arrow(ax, 10.0, 1.8, 7.5, 4.0, '검증', C_GRAY, lw=1.0)

    fig.tight_layout()
    return fig_to_bytes(fig)


def create_performance_reversal():
    """Performance reversal: Redis vs DB connection model"""
    fig, ax = plt.subplots(figsize=(9, 4.5))
    fig.set_facecolor(C_WHITE)

    # Data
    phases = [
        'Phase 6\n비관적 락',
        'Phase 10\nSQS 비동기',
        'Phase 10-2\nRedis 중복체크',
        'Phase 10-3\nPooling+Pipeline',
    ]
    values = [1730, 516, 1300, 658]
    colors = ['#94A3B8', C_ACCENT, C_DANGER, C_SUCCESS]

    bars = ax.bar(phases, values, color=colors, width=0.55, edgecolor='white', linewidth=2)

    for bar, val, col in zip(bars, values, colors):
        ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 30,
                f'{val:,}ms', ha='center', fontsize=12, fontweight='bold',
                color=col, fontfamily=FONT_KR)

    # Highlight the reversal
    ax.annotate('성능 역전!', xy=(2, 1300), xytext=(2.8, 1600),
                fontsize=13, fontweight='bold', color=C_DANGER,
                fontfamily=FONT_KR,
                arrowprops=dict(arrowstyle='->', color=C_DANGER, lw=2))

    # Explanation box
    ax.add_patch(FancyBboxPatch((2.55, 200), 1.5, 600, boxstyle="round,pad=0.1",
                                facecolor='#FEF2F2', edgecolor=C_DANGER, linewidth=1.5,
                                linestyle='--', alpha=0.7))

    ax.set_title('투표 API p95 응답시간 변화 추이', fontsize=14, fontweight='bold',
                 color=C_DARK, pad=15, fontfamily=FONT_KR)
    ax.set_ylabel('p95 (ms)', fontsize=11, color=C_GRAY, fontfamily=FONT_KR)
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_color('#E5E7EB')
    ax.spines['bottom'].set_color('#E5E7EB')
    ax.tick_params(colors=C_GRAY)
    for label in ax.get_xticklabels() + ax.get_yticklabels():
        label.set_fontfamily(FONT_KR)

    fig.tight_layout()
    return fig_to_bytes(fig)


def create_connection_model_diagram():
    """Redis single connection vs HikariCP pool comparison"""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4.5))
    fig.set_facecolor(C_WHITE)

    # ── Left: Lettuce single connection ──
    ax1.set_xlim(0, 10); ax1.set_ylim(0, 10)
    ax1.axis('off')
    ax1.set_title('Lettuce 기본 설정', fontsize=13, fontweight='bold',
                  color=C_DANGER, fontfamily=FONT_KR, pad=10)

    # Users
    for i in range(5):
        y = 8.5 - i * 1.2
        ax1.add_patch(FancyBboxPatch((0.3, y-0.2), 1.8, 0.5, boxstyle="round,pad=0.05",
                                    facecolor='#FEE2E2', edgecolor=C_DANGER, linewidth=1))
        ax1.text(1.2, y+0.05, f'요청 {i+1}', ha='center', va='center',
                 fontsize=8, color=C_DANGER, fontweight='bold', fontfamily=FONT_KR)
        # Arrow to single connection
        ax1.annotate('', xy=(4.5, 5.0), xytext=(2.1, y+0.05),
                     arrowprops=dict(arrowstyle='->', color='#FCA5A5', lw=1))

    # Single TCP connection - bottleneck
    ax1.add_patch(FancyBboxPatch((4.3, 4.0), 2.5, 2.0, boxstyle="round,pad=0.1",
                                facecolor=C_DANGER, edgecolor='#991B1B', linewidth=2))
    ax1.text(5.55, 5.3, 'TCP 1개', fontsize=12, fontweight='bold',
             color=C_WHITE, ha='center', fontfamily=FONT_KR)
    ax1.text(5.55, 4.7, '직렬화 병목', fontsize=9,
             color='#FCA5A5', ha='center', fontfamily=FONT_KR)
    ax1.text(5.55, 4.3, '50 → 1 → 50', fontsize=8,
             color='#FCA5A5', ha='center', fontfamily=FONT_KR)

    # Redis server
    ax1.add_patch(FancyBboxPatch((7.5, 4.3), 2.0, 1.5, boxstyle="round,pad=0.08",
                                facecolor='#7F1D1D', edgecolor='#991B1B', linewidth=1.5))
    ax1.text(8.5, 5.05, 'Redis', fontsize=11, fontweight='bold',
             color=C_WHITE, ha='center', fontfamily=FONT_KR)
    ax1.annotate('', xy=(7.5, 5.0), xytext=(6.8, 5.0),
                 arrowprops=dict(arrowstyle='->', color='#991B1B', lw=2))

    ax1.text(5, 1.5, 'p95: 1,300ms', fontsize=14, fontweight='bold',
             color=C_DANGER, ha='center', fontfamily=FONT_KR)

    # ── Right: HikariCP pool ──
    ax2.set_xlim(0, 10); ax2.set_ylim(0, 10)
    ax2.axis('off')
    ax2.set_title('HikariCP 커넥션 풀', fontsize=13, fontweight='bold',
                  color=C_SUCCESS, fontfamily=FONT_KR, pad=10)

    # Users
    for i in range(5):
        y = 8.5 - i * 1.2
        ax2.add_patch(FancyBboxPatch((0.3, y-0.2), 1.8, 0.5, boxstyle="round,pad=0.05",
                                    facecolor='#DCFCE7', edgecolor=C_SUCCESS, linewidth=1))
        ax2.text(1.2, y+0.05, f'요청 {i+1}', ha='center', va='center',
                 fontsize=8, color=C_SUCCESS, fontweight='bold', fontfamily=FONT_KR)

    # Multiple connections
    conn_y_positions = [7.5, 6.0, 4.5, 3.0]
    for i, cy in enumerate(conn_y_positions):
        ax2.add_patch(FancyBboxPatch((3.8, cy-0.2), 1.6, 0.5, boxstyle="round,pad=0.05",
                                    facecolor=C_SUCCESS, edgecolor='#065F46', linewidth=1))
        ax2.text(4.6, cy+0.05, f'Conn {i+1}', ha='center', va='center',
                 fontsize=7.5, color=C_WHITE, fontweight='bold', fontfamily=FONT_KR)
        ax2.annotate('', xy=(6.0, cy+0.05), xytext=(5.4, cy+0.05),
                     arrowprops=dict(arrowstyle='->', color=C_SUCCESS, lw=1.2))

    # Pool label
    ax2.add_patch(FancyBboxPatch((3.5, 2.0), 2.2, 6.5, boxstyle="round,pad=0.15",
                                facecolor='none', edgecolor=C_SUCCESS, linewidth=2,
                                linestyle='--'))
    ax2.text(4.6, 2.3, '30개 풀', fontsize=9, color=C_SUCCESS,
             ha='center', fontweight='bold', fontfamily=FONT_KR)

    # Arrows from users to connections
    for i in range(5):
        y_from = 8.5 - i * 1.2
        y_to = conn_y_positions[min(i, 3)]
        ax2.annotate('', xy=(3.8, y_to+0.05), xytext=(2.1, y_from+0.05),
                     arrowprops=dict(arrowstyle='->', color='#86EFAC', lw=1))

    # MySQL server
    ax2.add_patch(FancyBboxPatch((6.2, 3.8), 2.0, 2.5, boxstyle="round,pad=0.08",
                                facecolor='#065F46', edgecolor=C_SUCCESS, linewidth=1.5))
    ax2.text(7.2, 5.05, 'MySQL', fontsize=11, fontweight='bold',
             color=C_WHITE, ha='center', fontfamily=FONT_KR)
    ax2.text(7.2, 4.5, '병렬 처리', fontsize=9,
             color='#86EFAC', ha='center', fontfamily=FONT_KR)

    ax2.text(5, 1.5, 'p95: 516ms', fontsize=14, fontweight='bold',
             color=C_SUCCESS, ha='center', fontfamily=FONT_KR)

    fig.tight_layout(w_pad=2)
    return fig_to_bytes(fig)


def create_tech_decision_summary():
    """Technology decision summary radar/table visual"""
    fig, ax = plt.subplots(figsize=(9, 3.5))
    fig.set_facecolor(C_WHITE)
    ax.axis('off')

    decisions = [
        ('캐시', 'Redis\n(Global Cache)', 'Scale-out 시 일관성 보장\nLocal Cache는 인스턴스 간 불일치', C_DANGER),
        ('동시성', '비관적 락\n(SELECT FOR UPDATE)', 'Hot-spot 환경 100% 정합성\n낙관적 락은 Retry Storm 발생', C_ACCENT),
        ('MQ', 'SQS\n(Standard Queue)', '순서 불필요 + 멱등성 설계\n비용 $1/mo, Zero-Ops', C_WARNING),
        ('데이터', 'DB = Source of Truth\nRedis = Fast Display', 'Redis 장애 시 DB 복구 가능\n역방향은 불가', C_SUCCESS),
    ]

    for i, (category, choice, reason, color) in enumerate(decisions):
        y = 3.0 - i * 0.82
        # Category badge
        ax.add_patch(FancyBboxPatch((0.1, y-0.15), 1.0, 0.45, boxstyle="round,pad=0.06",
                                   facecolor=color, edgecolor=color))
        ax.text(0.6, y+0.07, category, ha='center', va='center', fontsize=9,
                fontweight='bold', color=C_WHITE, fontfamily=FONT_KR)
        # Choice
        ax.text(1.5, y+0.07, choice, ha='left', va='center', fontsize=9,
                fontweight='bold', color=C_DARK, fontfamily=FONT_KR)
        # Reason
        ax.text(4.8, y+0.07, reason, ha='left', va='center', fontsize=7.5,
                color=C_GRAY, fontfamily=FONT_KR)

        if i < len(decisions) - 1:
            ax.axhline(y=y-0.35, xmin=0.01, xmax=0.99, color='#E5E7EB', linewidth=0.5)

    ax.set_xlim(0, 10)
    ax.set_ylim(-0.5, 3.5)
    fig.tight_layout()
    return fig_to_bytes(fig)


# ═══════════════════════════════════════════
#  DOCX BUILDER
# ═══════════════════════════════════════════

def hex_to_rgb(hex_color):
    h = hex_color.lstrip('#')
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex.lstrip("#")}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", "4")}" w:space="0" '
            f'w:color="{val.get("color", "auto")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)


def add_styled_paragraph(doc, text, font_size=10, bold=False, color=C_DARK,
                         alignment=WD_ALIGN_PARAGRAPH.LEFT, space_before=0,
                         space_after=6, font_name='Apple SD Gothic Neo'):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = hex_to_rgb(color)
    run.font.name = font_name
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{font_name}"/>')
    rPr.append(rFonts)
    return p


def add_section_number(doc, number, title, subtitle=''):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(4)

    # Number badge
    run_num = p.add_run(f'  {number}  ')
    run_num.font.size = Pt(14)
    run_num.font.bold = True
    run_num.font.color.rgb = hex_to_rgb(C_WHITE)
    run_num.font.name = 'Apple SD Gothic Neo'
    # Add shading to the run
    rPr = run_num._element.get_or_add_rPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{C_ACCENT.lstrip("#")}"/>')
    rPr.append(shading)
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Apple SD Gothic Neo"/>')
    rPr.append(rFonts)

    run_space = p.add_run('  ')
    run_space.font.size = Pt(14)

    run_title = p.add_run(title)
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = hex_to_rgb(C_DARK)
    run_title.font.name = 'Apple SD Gothic Neo'
    rPr2 = run_title._element.get_or_add_rPr()
    rFonts2 = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Apple SD Gothic Neo"/>')
    rPr2.append(rFonts2)

    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(12)
        run_sub = p2.add_run(subtitle)
        run_sub.font.size = Pt(10)
        run_sub.font.color.rgb = hex_to_rgb(C_GRAY)
        run_sub.font.name = 'Apple SD Gothic Neo'
        rPr3 = run_sub._element.get_or_add_rPr()
        rFonts3 = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Apple SD Gothic Neo"/>')
        rPr3.append(rFonts3)


def add_label(doc, text, color=C_ACCENT):
    """Add a small colored label"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f'  {text}  ')
    run.font.size = Pt(9)
    run.font.bold = True
    run.font.color.rgb = hex_to_rgb(color)
    run.font.name = 'Apple SD Gothic Neo'
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Apple SD Gothic Neo"/>')
    rPr.append(rFonts)
    return p


def add_insight_box(doc, text):
    """Add a highlighted insight quote"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.5)

    # Add left border via paragraph formatting
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:left w:val="single" w:sz="18" w:space="8" w:color="{C_ACCENT.lstrip("#")}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    # Background shading
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="EFF6FF"/>')
    pPr.append(shd)

    run_icon = p.add_run('INSIGHT  ')
    run_icon.font.size = Pt(8)
    run_icon.font.bold = True
    run_icon.font.color.rgb = hex_to_rgb(C_ACCENT)
    run_icon.font.name = 'Apple SD Gothic Neo'

    run = p.add_run(text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = hex_to_rgb(C_PRIMARY)
    run.font.name = 'Apple SD Gothic Neo'
    rPr = run._element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Apple SD Gothic Neo"/>')
    rPr.append(rFonts)


def add_styled_table(doc, headers, rows, col_widths=None, header_color='1A2744'):
    """Create a clean styled table"""
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove default borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E7EB"/>'
        f'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        set_cell_shading(cell, f'#{header_color}')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = hex_to_rgb(C_WHITE)
        run.font.name = 'Apple SD Gothic Neo'
        rPr = run._element.get_or_add_rPr()
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Apple SD Gothic Neo"/>')
        rPr.append(rFonts)

    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r+1].cells[c]
            cell.text = ''
            if r % 2 == 1:
                set_cell_shading(cell, '#F8FAFC')
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            text = str(val)
            run = p.add_run(text)
            run.font.size = Pt(8.5)
            run.font.color.rgb = hex_to_rgb(C_DARK)
            run.font.name = 'Apple SD Gothic Neo'
            rPr = run._element.get_or_add_rPr()
            rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Apple SD Gothic Neo"/>')
            rPr.append(rFonts)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="E5E7EB"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def build_document():
    doc = Document()

    # ── Page Setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # ── Set default font ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Apple SD Gothic Neo'
    font.size = Pt(10)
    font.color.rgb = hex_to_rgb(C_DARK)
    rPr = style.element.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="Apple SD Gothic Neo"/>')
    rPr.append(rFonts)

    # ═══════════════════════════════════════
    #  HERO BANNER
    # ═══════════════════════════════════════
    hero_img = create_hero_banner()
    doc.add_picture(hero_img, width=Cm(17))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tech stack
    add_styled_paragraph(doc, '', font_size=4, space_after=2)
    p_tech = doc.add_paragraph()
    p_tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tech.paragraph_format.space_after = Pt(2)

    techs = ['Java 21', 'Spring Boot 3.5', 'MySQL 8.0', 'Redis 7', 'AWS SQS',
             'QueryDSL', 'Docker', 'Terraform', 'k6', 'Prometheus']
    for i, tech in enumerate(techs):
        run = p_tech.add_run(f' {tech} ')
        run.font.size = Pt(8)
        run.font.color.rgb = hex_to_rgb(C_ACCENT)
        run.font.name = 'Apple SD Gothic Neo'
        if i < len(techs) - 1:
            sep = p_tech.add_run('  ·  ')
            sep.font.size = Pt(8)
            sep.font.color.rgb = hex_to_rgb('#CBD5E1')

    add_styled_paragraph(doc,
        '반려동물 사진 콘테스트 플랫폼. 챌린지별 투표 → 실시간 랭킹 시스템.',
        font_size=9, color=C_GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=2, space_after=4)
    add_styled_paragraph(doc,
        '테스트 데이터: 10K 회원 · 100K 엔트리 · 1M 투표  |  인프라: EC2 t3.small · RDS db.t3.micro',
        font_size=8, color='#94A3B8', alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=0, space_after=6)

    add_divider(doc)

    # ═══════════════════════════════════════
    #  SELLING POINT 1: DB Query Optimization
    # ═══════════════════════════════════════
    add_section_number(doc, '01', 'DB 쿼리 최적화',
                       'Ranking API p95 응답시간 4.76s → 146ms (97% 감소)')

    add_label(doc, 'PROBLEM', C_DANGER)
    add_styled_paragraph(doc,
        'k6 부하 테스트(50 VUs)에서 랭킹 API p95 응답시간이 4.76초로 측정되었다. '
        'Grafana 대시보드와 p6spy 쿼리 로그를 분석한 결과 세 가지 병목을 확인했다.',
        font_size=9.5, space_before=4)

    add_styled_table(doc,
        ['병목', '원인', '영향'],
        [
            ['데이터 전송량', 'Java Stream.limit()로 애플리케이션에서 자름\nDB에서 전체 데이터를 전송', '불필요한 네트워크/메모리 사용'],
            ['N+1 쿼리', '연관 엔티티를 개별 SELECT로 조회\n요청당 21개 쿼리 발생', 'DB 커넥션 점유 시간 증가'],
            ['filesort', 'vote_count 정렬에 인덱스 없음\n매 요청마다 임시 테이블 정렬', 'CPU 사용률 증가'],
        ])

    add_label(doc, 'SOLUTION', C_SUCCESS)
    add_styled_paragraph(doc,
        '상위 레이어부터 순차 최적화 — 상위 병목을 먼저 제거해야 하위 레이어의 정확한 개선 효과를 측정할 수 있다.',
        font_size=9.5, bold=True, space_before=4)

    # Waterfall chart
    waterfall_img = create_optimization_waterfall()
    doc.add_picture(waterfall_img, width=Cm(15))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_styled_table(doc,
        ['단계', '최적화 내용', 'Before', 'After', '개선율'],
        [
            ['1단계', 'Spring Data JPA Pageable 적용\nDB LIMIT절로 데이터 전송량 제한', 'p95 4,760ms', 'p95 884ms', '81%↓'],
            ['2단계', 'QueryDSL Fetch Join\nN+1 문제 해결 (21 → 1 쿼리)', 'p95 884ms', 'p95 234ms', '73%↓'],
            ['3단계', '복합 인덱스 추가\n(challenge_id, vote_count DESC)', 'p95 234ms', 'p95 146ms', '38%↓'],
        ])

    # Query + Throughput comparison
    query_img = create_query_comparison()
    doc.add_picture(query_img, width=Cm(15))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_label(doc, 'RESULT', C_ACCENT)
    add_styled_table(doc,
        ['지표', '최적화 전 (50 VUs)', '최적화 후 (300 VUs)', '개선'],
        [
            ['랭킹 API p95', '4,760ms', '178ms', '97% 감소'],
            ['요청당 쿼리 수', '21개', '1개', '95% 감소'],
            ['전체 처리량', '10.9 RPS', '308 RPS', '28배 증가'],
            ['에러율', '-', '0%', '300 VUs 안정'],
        ])

    add_insight_box(doc,
        '최적화 순서가 중요하다. 1단계에서 데이터 전송량을 줄이지 않으면 '
        '2단계 Fetch Join의 효과가 과대/과소 측정될 수 있다. '
        '측정 가능한 최적화를 위해 상위 레이어부터 제거하는 원칙을 적용했다.')

    add_divider(doc)

    # ═══════════════════════════════════════
    #  SELLING POINT 2: Concurrency Control
    # ═══════════════════════════════════════
    add_section_number(doc, '02', '동시성 제어',
                       '3가지 전략 비교 분석 후 비관적 락 채택')

    add_label(doc, 'PROBLEM', C_DANGER)
    add_styled_paragraph(doc,
        '투표 기능에서 두 가지 레이스 컨디션이 발생했다.',
        font_size=9.5, space_before=4)

    add_styled_table(doc,
        ['레이스 컨디션', '발생 원인', '결과'],
        [
            ['Check-then-Act', 'SELECT(중복 확인)과 INSERT(투표 생성) 사이의 갭', '중복 투표 허용'],
            ['Lost Update', '동시에 voteCount를 읽고 +1하면 한쪽 유실', '투표 수 불일치'],
        ])

    add_styled_paragraph(doc,
        'synchronized는 JVM 레벨 락이라 다중 인스턴스에서 무효. '
        '@Transactional 단독으로는 MySQL REPEATABLE READ에서 Lost Update를 방지하지 못한다.',
        font_size=9, color=C_GRAY, space_before=4)

    add_label(doc, 'SOLUTION', C_SUCCESS)
    add_styled_paragraph(doc,
        '동일 조건(50명 동시 투표, 동일 Entry)에서 3가지 전략을 직접 구현하고 테스트했다.',
        font_size=9.5, space_before=4)

    concurrency_img = create_concurrency_comparison()
    doc.add_picture(concurrency_img, width=Cm(15))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_styled_table(doc,
        ['전략', '성공률', '소요시간', '문제점'],
        [
            ['비관적 락 (채택)\nSELECT FOR UPDATE', '100%', '1,730ms', '없음 — 순차 처리로 안정적'],
            ['낙관적 락\n@Version', '16%', '885ms', '42/50 실패. 기존 데이터 version NULL 이슈'],
            ['Atomic Update + Retry\n지수 백오프', '94%', '2,980ms', 'Retry Storm — 재시도가 재충돌 유발'],
        ])

    # Retry Storm diagram
    retry_img = create_retry_storm_diagram()
    doc.add_picture(retry_img, width=Cm(15))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_label(doc, 'TROUBLESHOOTING', C_WARNING)
    add_styled_table(doc,
        ['문제', '원인', '해결'],
        [
            ['@Version NULL', '기존 DB 데이터에 version 컬럼 값이 없음', 'ALTER TABLE로 기본값 마이그레이션'],
            ['@Transactional 미적용', '같은 클래스 내부 메서드 호출 시\nSpring AOP 프록시 우회', '@Lazy 자기 주입으로 프록시 통한 호출 보장'],
        ])

    add_insight_box(doc,
        '"느린 비관적 락이 역설적으로 빠르다." '
        '고경쟁(hot-spot) 환경에서는 낙관적 접근이 충돌 → 재시도 → 재충돌의 '
        'Retry Storm을 일으켜 전체 응답시간이 오히려 증가한다. '
        '경쟁 강도에 따라 전략을 선택해야 한다.')

    add_divider(doc)

    # ═══════════════════════════════════════
    #  SELLING POINT 3: Async Architecture + Performance Reversal
    # ═══════════════════════════════════════
    add_section_number(doc, '03', '비동기 투표 아키텍처',
                       '"빠른 기술 ≠ 빠른 시스템" — 커넥션 모델의 중요성')

    add_label(doc, 'PROBLEM', C_DANGER)
    add_styled_paragraph(doc,
        '비관적 락으로 정합성은 확보했지만 p95 1.73초로 사용자 체감 성능이 부족했다. '
        '동기 DB 쓰기가 병목.',
        font_size=9.5, space_before=4)

    add_label(doc, 'SOLUTION — 읽기/쓰기 분리 (CQRS)', C_SUCCESS)

    arch_img = create_async_architecture()
    doc.add_picture(arch_img, width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_styled_paragraph(doc,
        'Fast Path: 투표 → Redis ZINCRBY 즉시 반영 → 사용자에게 즉시 응답 (~5ms)\n'
        'Slow Path: SQS 메시지 발행 → Consumer가 비동기로 DB 영속화\n'
        '정합성: 5분 주기 Scheduler가 Redis ↔ DB 정합성 검증 및 보정',
        font_size=9.5, space_before=4)

    add_label(doc, 'MQ 선택 — SQS를 선택한 이유', C_ACCENT)
    add_styled_table(doc,
        ['기준', 'Kafka (MSK)', 'RabbitMQ', 'SQS (채택)'],
        [
            ['월 비용', '$200+', '$50-100', '$1'],
            ['운영 부담', 'ZooKeeper 관리', 'Erlang 관리', 'Zero-Ops (완전 관리형)'],
            ['확장', '수동 파티션', '수동', '자동 (무제한)'],
            ['적합성', '이벤트 스트리밍에 적합', '복잡한 라우팅에 적합', '단순 큐잉에 최적'],
        ])

    add_styled_paragraph(doc,
        '투표 메시지는 순서 보장이 불필요하고, 멱등성 설계(Application 중복 체크 + DB Unique Constraint)로 '
        '중복 처리가 가능하므로 SQS Standard Queue를 선택했다. '
        'DLQ 설정(3회 재시도 → 14일 보관)으로 메시지 유실을 방지한다.',
        font_size=9, color=C_GRAY, space_before=4)

    add_label(doc, 'DISCOVERY — 성능 역전 현상', '#7C3AED')
    add_styled_paragraph(doc,
        '중복 투표 체크를 DB에서 Redis(SISMEMBER)로 이관했더니 오히려 2.5배 느려졌다.',
        font_size=9.5, bold=True, space_before=4, color=C_DANGER)

    reversal_img = create_performance_reversal()
    doc.add_picture(reversal_img, width=Cm(15))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_label(doc, 'ROOT CAUSE — 커넥션 모델의 차이', C_DANGER)

    conn_img = create_connection_model_diagram()
    doc.add_picture(conn_img, width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_styled_table(doc,
        ['', 'Lettuce (Redis)', 'HikariCP (MySQL)'],
        [
            ['커넥션', '단일 TCP 멀티플렉싱', '30개 커넥션 풀'],
            ['50 동시 요청', '1개 파이프에 직렬화 → 병목', '30개로 병렬 분산 → 자연 로드밸런싱'],
            ['결과', 'p95 1,300ms', 'p95 516ms'],
        ])

    add_label(doc, 'RESOLUTION', C_SUCCESS)
    add_styled_paragraph(doc,
        'Lettuce Connection Pooling(commons-pool2, max-active: 20) + '
        'Redis Pipeline(ZINCRBY + ZSCORE를 단일 RTT로) 적용. '
        '최종적으로 DB 중복 체크를 유지하는 Hybrid 전략을 채택했다.',
        font_size=9.5, space_before=4)

    add_label(doc, 'RESULT', C_ACCENT)

    add_styled_table(doc,
        ['단계', 'p95', '설명'],
        [
            ['Phase 6 — 비관적 락', '1,730ms', '동기 DB 쓰기 병목'],
            ['Phase 10 — SQS 비동기', '516ms', 'DB 쓰기 분리 (70%↓)'],
            ['Phase 10-2 — Redis 중복체크', '1,300ms', '성능 역전 (Lettuce 단일 커넥션 병목)'],
            ['Phase 10-3 — Pooling + Pipeline', '658ms', '최종 (비관적 락 대비 62%↓)'],
        ])

    add_insight_box(doc,
        '"빠른 기술 ≠ 빠른 시스템." '
        'Redis의 O(1) 연산 속도보다 HikariCP 30개 커넥션의 병렬 처리가 '
        '50 동시 사용자 환경에서 더 효과적이었다. '
        '개별 연산 속도가 아닌 커넥션 모델과 동시성 처리 방식이 실제 성능을 결정한다.')

    add_divider(doc)

    # ═══════════════════════════════════════
    #  TECH DECISIONS SUMMARY
    # ═══════════════════════════════════════
    add_styled_paragraph(doc, '기술적 의사결정 요약', font_size=14, bold=True,
                         color=C_DARK, space_before=16, space_after=8)

    tech_img = create_tech_decision_summary()
    doc.add_picture(tech_img, width=Cm(16))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Save ──
    output_path = os.path.join(OUT_DIR, 'PetStar_Portfolio_v1.docx')
    doc.save(output_path)
    print(f'Portfolio saved to: {output_path}')
    return output_path


if __name__ == '__main__':
    path = build_document()
    # Cleanup temp images
    import shutil
    if os.path.exists(IMG_DIR):
        shutil.rmtree(IMG_DIR)
    print('Done!')
